"""
学术综述格式检查和文档生成模块

此模块负责：
1. 判断论文语言（中文/英文）
2. 生成格式化的参考文献（IEEE格式和GB/T 7714-2015格式）
3. 将综述保存为格式化的Word文档
4. 处理文档格式、字体、标题编号等
"""

import re
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def is_chinese_paper(paper):
    """
    判断论文是否为中文论文

    Args:
        paper: 论文信息字典

    Returns:
        True if 中文论文, False if 英文论文
    """
    # 检查标题
    title = paper.get('title', '')
    if title:
        # 检查是否包含中文字符
        if re.search(r'[\u4e00-\u9fff]', title):
            return True

    # 检查摘要
    summary = paper.get('summary', '')
    if summary:
        if re.search(r'[\u4e00-\u9fff]', summary):
            return True

    # 检查venue（期刊/会议名称）
    venue = paper.get('venue', '')
    if venue:
        # 常见中文期刊关键词
        chinese_venue_keywords = ['学报', '期刊', '杂志', '研究', '科学', '工程', '技术', '信息', '计算机', '电子', '通信']
        venue_lower = venue.lower()
        for keyword in chinese_venue_keywords:
            if keyword in venue:
                # 进一步检查是否包含中文字符
                if re.search(r'[\u4e00-\u9fff]', venue):
                    return True

    # 默认判断为英文论文
    return False


def format_reference_ieee(paper, index):
    """
    将论文信息格式化为IEEE格式的参考文献

    Args:
        paper: 论文信息字典
        index: 参考文献序号

    Returns:
        格式化的参考文献字符串
    """
    title = paper.get('title', '未知标题')
    authors = paper.get('authors', [])
    published = paper.get('published', '')
    venue = paper.get('venue', '')
    source = paper.get('source', 'unknown')
    entry_id = paper.get('entry_id', '')

    # 格式化作者（IEEE格式：First Last, First Last, and First Last）
    if authors:
        if len(authors) == 1:
            author_str = authors[0]
        elif len(authors) <= 6:
            # 前N-1个用逗号分隔，最后一个用"and"连接
            author_str = ', '.join(authors[:-1]) + ', and ' + authors[-1]
        else:
            # 超过6个作者，只列出前6个，然后加"et al."
            author_str = ', '.join(authors[:6]) + ', et al.'
    else:
        author_str = 'Anonymous'

    # 提取年份
    year = ''
    if published:
        year_match = re.search(r'(\d{4})', published)
        if year_match:
            year = year_match.group(1)

    # 判断文献类型并格式化（IEEE标准）
    if source.lower() == 'arxiv' or 'arxiv' in venue.lower() or 'arxiv' in title.lower():
        # 预印本格式：[序号] Author(s). "Title," arXiv preprint arXiv:XXXX.XXXX, Year.
        arxiv_id = ''
        if entry_id:
            arxiv_match = re.search(r'arxiv\.org/abs/([\d.]+)', entry_id)
            if arxiv_match:
                arxiv_id = arxiv_match.group(1)
        if arxiv_id:
            ref_str = f"[{index}] {author_str}, \"{title},\" arXiv preprint arXiv:{arxiv_id}"
        else:
            ref_str = f"[{index}] {author_str}, \"{title},\" arXiv preprint"
        if year:
            ref_str += f", {year}"
        ref_str += "."
    elif venue:
        # 判断是期刊还是会议
        venue_lower = venue.lower()
        if any(keyword in venue_lower for keyword in ['conference', 'proceedings', 'workshop', 'symposium', 'meeting']):
            # 会议论文格式：[序号] Author(s), "Title," in Proc. Conference Name, Location, Year, pp. pages.
            ref_str = f"[{index}] {author_str}, \"{title},\" in Proc. {venue}"
            if year:
                ref_str += f", {year}"
            ref_str += "."
        else:
            # 期刊论文格式：[序号] Author(s), "Title," Journal Name, vol. vol, no. no, pp. pages, Month Year.
            ref_str = f"[{index}] {author_str}, \"{title},\" {venue}"
            if year:
                ref_str += f", {year}"
            ref_str += "."
    else:
        # 其他类型
        ref_str = f"[{index}] {author_str}, \"{title}\""
        if year:
            ref_str += f", {year}"
        if venue:
            ref_str += f", {venue}"
        ref_str += "."

    return ref_str


def format_reference_gb7714(paper, index):
    """
    将论文信息格式化为GB/T 7714-2015格式的参考文献

    Args:
        paper: 论文信息字典
        index: 参考文献序号

    Returns:
        格式化的参考文献字符串
    """
    title = paper.get('title', '未知标题')
    authors = paper.get('authors', [])
    published = paper.get('published', '')
    venue = paper.get('venue', '')
    source = paper.get('source', 'unknown')
    entry_id = paper.get('entry_id', '')

    # 格式化作者（最多3个，超过用"等"）
    if authors:
        if len(authors) <= 3:
            author_str = ', '.join(authors)
        else:
            author_str = ', '.join(authors[:3]) + ', 等'
    else:
        author_str = '佚名'

    # 提取年份
    year = ''
    if published:
        # 尝试从日期中提取年份
        year_match = re.search(r'(\d{4})', published)
        if year_match:
            year = year_match.group(1)

    # 判断文献类型并格式化（GB/T 7714-2015标准）
    if source.lower() == 'arxiv' or 'arxiv' in venue.lower() or 'arxiv' in title.lower():
        # 预印本格式：[序号] 作者. 题名[EB/OL]. (发表日期)[引用日期]. 获取和访问路径.
        date_str = published if published else '未知日期'
        ref_str = f"[{index}] {author_str}. {title}[EB/OL]. ({date_str})"
        if entry_id:
            ref_str += f"[{datetime.now().strftime('%Y-%m-%d')}]. {entry_id}."
        else:
            ref_str += "."
    elif venue:
        # 判断是期刊还是会议
        venue_lower = venue.lower()
        if any(keyword in venue_lower for keyword in ['conference', 'proceedings', 'workshop', 'symposium', 'meeting']):
            # 会议论文格式：[序号] 作者. 题名[C]//会议名称. 会议地点: 出版者, 出版年: 起止页码.
            ref_str = f"[{index}] {author_str}. {title}[C]//{venue}"
            if year:
                ref_str += f". {year}"
            ref_str += "."
        else:
            # 期刊论文格式：[序号] 作者. 题名[J]. 刊名, 出版年, 卷(期): 起止页码.
            ref_str = f"[{index}] {author_str}. {title}[J]. {venue}"
            if year:
                ref_str += f", {year}"
            ref_str += "."
    else:
        # 其他类型，使用通用格式：[序号] 作者. 题名[M/D/R等]. 出版地: 出版者, 出版年.
        ref_str = f"[{index}] {author_str}. {title}"
        if year:
            ref_str += f". {year}"
        if venue:
            ref_str += f". {venue}"
        ref_str += "."

    return ref_str


def save_review_to_docx(keyword, papers, unique_final_papers, review, filename):
    """
    将综述保存为格式化的docx文件

    Args:
        keyword: 关键词
        papers: 所有搜索到的论文
        unique_final_papers: 最终选择的论文
        review: 生成的综述内容
        filename: 保存的文件名
    """
    doc = Document()

    # 设置字体函数：中文使用仿宋，英文使用Times New Roman
    def set_font(run, chinese_font='仿宋', english_font='Times New Roman', size=12):
        """
        设置字体：中文使用仿宋，英文、数字使用Times New Roman
        """
        # 设置英文字体
        run.font.name = english_font
        # 设置中文字体（东亚字体）
        run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
        run.font.size = Pt(size)

    # 设置文本字体：中文用仿宋，数字和字母用Times New Roman
    def set_text_font(para, text, chinese_font='仿宋', english_font='Times New Roman', size=12):
        """
        将文本按字符分割，中文用仿宋，数字和字母用Times New Roman
        """
        # 匹配中文字符、数字字母、标点符号
        pattern = r'([\u4e00-\u9fff]+|[a-zA-Z0-9]+|[^\u4e00-\u9fff\w\s]+|\s+)'
        parts = re.findall(pattern, text)

        for part in parts:
            if not part.strip():
                # 空白字符
                run = para.add_run(part)
                set_font(run, chinese_font=chinese_font, english_font=english_font, size=size)
            elif re.match(r'^[a-zA-Z0-9]+$', part):
                # 纯数字和字母，使用Times New Roman
                run = para.add_run(part)
                run.font.name = english_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)  # 保持中文字体设置
                run.font.size = Pt(size)
            elif re.match(r'^[\u4e00-\u9fff]+$', part):
                # 纯中文，使用仿宋
                run = para.add_run(part)
                set_font(run, chinese_font=chinese_font, english_font=english_font, size=size)
            else:
                # 混合内容，需要逐字符处理
                for char in part:
                    if re.match(r'^[a-zA-Z0-9]$', char):
                        # 数字和字母
                        run = para.add_run(char)
                        run.font.name = english_font
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                        run.font.size = Pt(size)
                    else:
                        # 中文或其他字符
                        run = para.add_run(char)
                        set_font(run, chinese_font=chinese_font, english_font=english_font, size=size)

    # 兼容旧函数名
    def set_chinese_font(run, font_name='仿宋', size=12):
        set_font(run, chinese_font=font_name, size=size)

    # 标题（仿宋小三，15磅）
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_text_font(title_para, f'学术综述：{keyword}', chinese_font='仿宋', english_font='Times New Roman', size=15)
    for run in title_para.runs:
        run.font.bold = True  # 加粗

    # 添加空行
    doc.add_paragraph()

    # 基本信息（一级标题格式：1.）
    info_heading_para = doc.add_paragraph()
    set_text_font(info_heading_para, '1. 基本信息', chinese_font='仿宋', english_font='Times New Roman', size=14)
    for run in info_heading_para.runs:
        run.font.bold = True
    info_heading_para.paragraph_format.space_after = Pt(12)
    info_heading_para.paragraph_format.space_before = Pt(12)

    info_para = doc.add_paragraph()
    # 添加"关键词："部分
    run1 = info_para.add_run('关键词：')
    run1.bold = True
    set_font(run1, chinese_font='仿宋', english_font='Times New Roman', size=12)
    # 添加关键词内容（使用set_text_font处理数字和字母）
    for char in keyword:
        if re.match(r'^[a-zA-Z0-9]$', char):
            # 数字和字母用Times New Roman
            run = info_para.add_run(char)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            run.font.size = Pt(12)
        else:
            # 中文或其他字符用仿宋
            run = info_para.add_run(char)
            set_font(run, chinese_font='仿宋', english_font='Times New Roman', size=12)

    info_para = doc.add_paragraph()
    run1 = info_para.add_run('搜索到的论文数量：')
    run1.bold = True
    set_font(run1, chinese_font='仿宋', english_font='Times New Roman', size=12)
    set_text_font(info_para, f'{len(papers)} 篇（去重后）', chinese_font='仿宋', english_font='Times New Roman', size=12)

    info_para = doc.add_paragraph()
    run1 = info_para.add_run('最终选择的论文数量：')
    run1.bold = True
    set_font(run1, chinese_font='仿宋', english_font='Times New Roman', size=12)
    set_text_font(info_para, f'{len(unique_final_papers)} 篇', chinese_font='仿宋', english_font='Times New Roman', size=12)

    # 统计来源
    sources = {}
    for paper in unique_final_papers:
        source = paper.get('source', 'unknown')
        sources[source] = sources.get(source, 0) + 1
    if sources:
        info_para = doc.add_paragraph()
        run1 = info_para.add_run('数据来源：')
        run1.bold = True
        set_font(run1, chinese_font='仿宋', english_font='Times New Roman', size=12)
        set_text_font(info_para, ', '.join([f'{k}({v}篇)' for k, v in sources.items()]), chinese_font='仿宋', english_font='Times New Roman', size=12)

    # 添加生成时间
    info_para = doc.add_paragraph()
    run1 = info_para.add_run('生成时间：')
    run1.bold = True
    set_font(run1, chinese_font='仿宋', english_font='Times New Roman', size=12)
    set_text_font(info_para, datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'), chinese_font='仿宋', english_font='Times New Roman', size=12)

    # 添加分页符
    doc.add_page_break()

    # 处理综述内容，解析Markdown格式的标题和段落
    review_lines = review.split('\n')

    # 查找"参考文献"部分的位置，如果存在则跳过（我们会在文档末尾单独生成）
    ref_section_start = -1
    for idx, line in enumerate(review_lines):
        stripped = line.strip()
        if stripped.startswith('##') and ('参考文献' in stripped or '参考' in stripped and '文献' in stripped):
            ref_section_start = idx
            break

    # 如果找到参考文献部分，只处理之前的内容
    if ref_section_start >= 0:
        review_lines = review_lines[:ref_section_start]

    # 标题编号计数器
    level1_counter = 0  # 引言后的二级标题编号（1, 2, 3...）
    level2_counter = 0  # 三级标题编号（在每个二级标题下从1开始：1.1, 1.2... 2.1, 2.2...）
    level3_counter = 0  # 四级标题编号（如果需要：1.1.1, 1.1.2...）
    last_level = 0  # 上一个标题的级别
    numbering_started = False  # 标记是否开始编号（从引言开始）
    after_intro = False  # 标记是否在引言之后（引言下不需要其他次级标题）
    first_level1_title = True  # 标记是否是第一个一级标题（总标题，居中无序号）

    # 定义不需要编号的标题关键词
    no_numbering_keywords = ['摘要', '关键词']

    # 辅助函数：去除标题中已有的编号
    def remove_existing_numbering(text):
        """去除标题开头的编号（如"1 "、"2 "、"1.1 "、"1.1.1 "、"0 "等）"""
        import re
        # 匹配开头的编号模式：
        # - "0 "（引言）
        # - "1 "、"2 "、"3 "等（引言后的二级标题）
        # - "1.1 "、"1.2 "、"2.1 "等（三级标题）
        # - "1.1.1 "、"1.1.2 "等（四级标题）
        # 确保编号后必须有空格，避免误匹配标题中的数字
        pattern = r'^(0\s+|\d+(\.\d+)*(\.\d+)?\s+)'
        return re.sub(pattern, '', text).strip()

    i = 0
    while i < len(review_lines):
        line = review_lines[i].strip()

        # 跳过空行（段落之间不要空行）
        if not line:
            i += 1
            continue

        # 检查是否是标题
        if line.startswith('#### '):
            # 四级标题（仿宋小四，12磅，格式：1.1.1, 1.1.2... 1.2.1, 1.2.2...）
            title_text = line[5:].strip()
            # 去除标题中可能已有的编号
            title_text = remove_existing_numbering(title_text)

            # 如果还没有开始编号（引言之前），则不编号
            if not numbering_started:
                numbered_title = title_text
            # 如果是在引言之后（引言下不需要其他次级标题），跳过
            elif after_intro:
                # 跳过这个标题，不处理
                pass
            else:
                # 引言之后的二级标题下的三级标题下的四级标题，格式：1.1.1, 1.1.2... 1.2.1, 1.2.2...
                # 更新计数器：如果上一个标题是四级，则递增；否则重置为1
                if last_level == 4:
                    level3_counter += 1
                else:
                    level3_counter = 1

                # 生成编号：{level1_counter}.{level2_counter}.{level3_counter}
                number = f"{level1_counter}.{level2_counter}.{level3_counter} "
                numbered_title = number + title_text

                heading_para = doc.add_paragraph()
                set_text_font(heading_para, numbered_title, chinese_font='仿宋', english_font='Times New Roman', size=12)
                for run in heading_para.runs:
                    run.font.bold = True
                heading_para.paragraph_format.space_after = Pt(12)
                heading_para.paragraph_format.space_before = Pt(12)

                last_level = 4
        elif line.startswith('### '):
            # 三级标题（仿宋小四，12磅，格式：1.1, 1.2... 2.1, 2.2...）
            title_text = line[4:].strip()
            # 去除标题中可能已有的编号
            title_text = remove_existing_numbering(title_text)

            # 如果还没有开始编号（引言之前），则不编号
            if not numbering_started:
                numbered_title = title_text
            # 如果是在引言之后（引言下不需要其他次级标题），跳过
            elif after_intro:
                # 跳过这个标题，不处理
                pass
            else:
                # 引言之后的二级标题下的三级标题，格式：1.1, 1.2... 2.1, 2.2...
                # 更新计数器：如果上一个标题是三级，则递增；否则重置为1
                if last_level == 3:
                    level2_counter += 1
                else:
                    level2_counter = 1
                level3_counter = 0  # 重置四级标题计数器

                # 生成编号：{level1_counter}.{level2_counter}
                number = f"{level1_counter}.{level2_counter} "
                numbered_title = number + title_text

                heading_para = doc.add_paragraph()
                set_text_font(heading_para, numbered_title, chinese_font='仿宋', english_font='Times New Roman', size=12)
                for run in heading_para.runs:
                    run.font.bold = True
                heading_para.paragraph_format.space_after = Pt(12)
                heading_para.paragraph_format.space_before = Pt(12)

                last_level = 3
        elif line.startswith('## '):
            # 二级标题（仿宋小四，12磅）
            title_text = line[3:].strip()
            # 去除标题中可能已有的编号（但保留"0 "用于引言的特殊处理）
            # 先检查是否是引言，如果是引言且已有"0 "，则保留；否则去除编号
            if '引言' not in title_text:
                title_text = remove_existing_numbering(title_text)
            elif title_text.startswith('0 '):
                # 引言已有"0 "编号，去除它（后面会重新添加）
                title_text = title_text[2:].strip()

            # 检查是否是引言（开始编号）
            if '引言' in title_text and not numbering_started:
                numbering_started = True
                after_intro = True  # 标记在引言之后（引言下不需要其他次级标题）
                level1_counter = 1  # 引言后的二级标题从2开始（引言算作第1章，但不显示编号）
                level2_counter = 0  # 三级标题计数器
                level3_counter = 0
                numbered_title = title_text  # 引言不显示编号
                last_level = 2  # 引言是二级标题
            # 检查是否是不需要编号的标题（摘要、关键词）
            elif any(keyword in title_text for keyword in no_numbering_keywords):
                numbered_title = title_text
                last_level = 2
            # 如果还没有开始编号（引言之前），则不编号
            elif not numbering_started:
                numbered_title = title_text
                last_level = 2
            # 如果是在引言之后，遇到第一个二级标题，清除after_intro标志，从1开始编号
            elif after_intro:
                after_intro = False  # 清除引言后标志
                level1_counter = 1  # 第一个二级标题从1开始
                level2_counter = 0  # 重置三级标题计数器
                level3_counter = 0  # 重置四级标题计数器
                number = f"{level1_counter} "
                numbered_title = number + title_text
                last_level = 2
            else:
                # 引言之后的后续二级标题，递增编号（2, 3, 4...）
                level1_counter += 1
                level2_counter = 0  # 重置三级标题计数器
                level3_counter = 0  # 重置四级标题计数器

                # 生成编号（格式：2, 3, 4等）
                number = f"{level1_counter} "
                numbered_title = number + title_text
                last_level = 2

            heading_para = doc.add_paragraph()
            set_text_font(heading_para, numbered_title, chinese_font='仿宋', english_font='Times New Roman', size=12)
            for run in heading_para.runs:
                run.font.bold = True
            heading_para.paragraph_format.space_after = Pt(12)
            heading_para.paragraph_format.space_before = Pt(12)
        elif line.startswith('# '):
            # 一级标题（仿宋四号，14磅）
            title_text = line[2:].strip()

            # 第一个一级标题是总标题，居中，不需要数字序号
            if first_level1_title:
                numbered_title = title_text
                heading_para = doc.add_paragraph()
                heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
                set_text_font(heading_para, numbered_title, chinese_font='仿宋', english_font='Times New Roman', size=14)
                for run in heading_para.runs:
                    run.font.bold = True
                heading_para.paragraph_format.space_after = Pt(12)
                heading_para.paragraph_format.space_before = Pt(12)
                first_level1_title = False
                last_level = 1
            # 其他一级标题按1. 2. 等编号（但根据新要求，一级标题应该很少出现）
            else:
                # 检查是否是不需要编号的标题
                if any(keyword in title_text for keyword in no_numbering_keywords):
                    numbered_title = title_text
                else:
                    # 更新计数器
                    level1_counter += 1
                    level2_counter = 0  # 重置二级标题计数器
                    level3_counter = 0  # 重置三级标题计数器

                    # 生成编号
                    number = f"{level1_counter}. "
                    numbered_title = number + title_text

                heading_para = doc.add_paragraph()
                set_text_font(heading_para, numbered_title, chinese_font='仿宋', english_font='Times New Roman', size=14)
                for run in heading_para.runs:
                    run.font.bold = True
                heading_para.paragraph_format.space_after = Pt(12)
                heading_para.paragraph_format.space_before = Pt(12)
                last_level = 1
        else:
            # 普通段落（中文用仿宋，数字和字母用Times New Roman）
            current_para = doc.add_paragraph()

            # 处理文本中的参考文献引用 [1] 或 [1,2,3]，转换为上标
            text = line

            # 去掉类似*的不必要标点符号（保留必要的标点符号）
            # 去掉单独的*号（前后有空格或开头结尾的*）
            text = re.sub(r'\s*\*\s*', ' ', text)  # 去掉单独的*号及其周围空格
            text = re.sub(r'^\*+', '', text)  # 去掉开头的*号
            text = re.sub(r'\*+$', '', text)  # 去掉结尾的*号
            text = re.sub(r'\*{2,}', '', text)  # 去掉连续的多个*号
            # 清理多余的空格
            text = re.sub(r'\s+', ' ', text).strip()

            # 使用正则表达式匹配 [数字] 或 [数字,数字,...]
            ref_pattern = r'\[(\d+(?:,\d+)*)\]'

            last_end = 0
            for match in re.finditer(ref_pattern, text):
                # 添加引用前的文本（使用set_text_font处理数字和字母）
                if match.start() > last_end:
                    prefix_text = text[last_end:match.start()]
                    set_text_font(current_para, prefix_text, chinese_font='仿宋', english_font='Times New Roman', size=12)

                # 添加引用（上标格式，数字用Times New Roman）
                ref_text = '[' + match.group(1) + ']'
                # 引用中的数字和括号需要分别处理
                for char in ref_text:
                    run = current_para.add_run(char)
                    if char.isdigit():
                        # 数字用Times New Roman
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                    else:
                        # 括号用仿宋
                        set_font(run, chinese_font='仿宋', english_font='Times New Roman', size=12)
                    run.font.size = Pt(12)
                    run.font.superscript = True  # 设置为上标

                last_end = match.end()

            # 添加剩余文本（使用set_text_font处理数字和字母）
            if last_end < len(text):
                suffix_text = text[last_end:]
                set_text_font(current_para, suffix_text, chinese_font='仿宋', english_font='Times New Roman', size=12)

            # 如果段落为空（只有引用），至少添加一个run
            if not current_para.runs:
                set_text_font(current_para, text, chinese_font='仿宋', english_font='Times New Roman', size=12)

            # 设置段落格式
            current_para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2个字符（约24磅）
            current_para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            current_para.paragraph_format.space_after = Pt(0)  # 段间距0磅（段落之间不要空行）

        i += 1

    # 添加分页符
    doc.add_page_break()

    # 参考文献列表（GB/T 7714-2015格式）
    ref_heading_para = doc.add_paragraph()
    ref_heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_text_font(ref_heading_para, '参考文献', chinese_font='仿宋', english_font='Times New Roman', size=14)
    for run in ref_heading_para.runs:
        run.font.bold = True  # 加粗
    ref_heading_para.paragraph_format.space_after = Pt(12)

    # 根据论文语言选择引用格式：中文期刊使用GB/T 7714-2015，英文期刊使用IEEE格式
    for i, paper in enumerate(unique_final_papers, 1):
        # 判断论文语言
        if is_chinese_paper(paper):
            ref_text = format_reference_gb7714(paper, i)
            # GB/T 7714-2015格式：悬挂缩进
            ref_para = doc.add_paragraph()
            set_text_font(ref_para, ref_text, chinese_font='仿宋', english_font='Times New Roman', size=12)
            ref_para.paragraph_format.left_indent = Pt(24)  # 整个段落左缩进
            ref_para.paragraph_format.first_line_indent = Pt(-24)  # 首行悬挂缩进（负值表示向左）
            ref_para.paragraph_format.space_after = Pt(6)  # 段后间距
            ref_para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
        else:
            ref_text = format_reference_ieee(paper, i)
            # IEEE格式：左对齐，首行缩进
            ref_para = doc.add_paragraph()
            set_text_font(ref_para, ref_text, chinese_font='仿宋', english_font='Times New Roman', size=12)
            ref_para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
            ref_para.paragraph_format.space_after = Pt(6)  # 段后间距
            ref_para.paragraph_format.line_spacing = 1.5  # 1.5倍行距

    # 保存文档
    import os
    import time

    # 检查文件是否已存在
    if os.path.exists(filename):
        print(f"\n警告：文件 '{filename}' 已存在")
        print("可能的原因：")
        print("  1. 文件正在被其他程序（如 Word）打开")
        print("  2. 之前已生成过同名文件")
        print("\n请关闭可能打开该文件的程序，然后重试")

        # 尝试生成新文件名
        base_name = os.path.splitext(filename)[0]
        extension = os.path.splitext(filename)[1]
        counter = 1
        new_filename = f"{base_name}_{counter}{extension}"
        while os.path.exists(new_filename):
            counter += 1
            new_filename = f"{base_name}_{counter}{extension}"

        print(f"\n或者使用新文件名保存: {new_filename}")
        use_new = input("是否使用新文件名保存？(y/n，默认y): ").strip().lower()
        if use_new in ['', 'y', 'yes', '是']:
            filename = new_filename
        else:
            print("取消保存操作")
            return

    try:
        # 尝试保存文件
        doc.save(filename)
        print(f"\n✓ 综述已成功保存到文件: {filename}")
        print(f"完整路径: {os.path.abspath(filename)}")
    except PermissionError as e:
        print(f"\n✗ 保存失败：权限错误")
        print(f"错误信息: {e}")
        print("\n可能的原因：")
        print("  1. 文件正在被其他程序（如 Microsoft Word）打开")
        print("  2. 文件被设置为只读")
        print("  3. 没有写入权限")
        print("\n解决方案：")
        print("  1. 关闭可能打开该文件的所有程序")
        print("  2. 检查文件属性，确保不是只读")
        print("  3. 检查文件夹的写入权限")
        print("  4. 尝试使用不同的文件名")

        # 提供重试选项
        retry = input("\n是否尝试使用新文件名保存？(y/n): ").strip().lower()
        if retry in ['y', 'yes', '是']:
            base_name = os.path.splitext(filename)[0]
            extension = os.path.splitext(filename)[1]
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            new_filename = f"{base_name}_{timestamp}{extension}"
            try:
                doc.save(new_filename)
                print(f"\n✓ 综述已成功保存到新文件: {new_filename}")
                print(f"完整路径: {os.path.abspath(new_filename)}")
            except Exception as e2:
                print(f"\n✗ 使用新文件名保存也失败: {e2}")
    except Exception as e:
        print(f"\n✗ 保存文件时发生错误: {e}")
        print(f"错误类型: {type(e).__name__}")
        print("\n请检查：")
        print("  1. 文件路径是否有效")
        print("  2. 是否有足够的磁盘空间")
        print("  3. 文件夹是否存在")

