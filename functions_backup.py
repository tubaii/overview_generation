from openai import OpenAI
import arxiv
import requests
import tarfile
import io
import re
import hashlib
import time
from typing import List, Dict, Set, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 初始化OpenAI客户端（使用ModelScope接口）
client = OpenAI(
    base_url='https://api-inference.modelscope.cn/v1',
    api_key='ms-900e2843-b140-4758-82e7-59c6358b737f', # ModelScope Token
)

def make_semantic_scholar_request(url, params=None, max_retries=15, retry_delay=2):
    """
    统一的Semantic Scholar API请求函数，包含重试和速率限制处理
    
    Args:
        url: API请求URL
        params: 请求参数
        max_retries: 最大重试次数（默认15次，增加以应对速率限制）
        retry_delay: 基础重试延迟（秒），会指数增长
    
    Returns:
        response对象或None（如果失败）
    """
    import requests.exceptions
    
    for attempt in range(max_retries):
        try:
            response = requests.get(url, params=params, timeout=30)
            
            # 成功响应
            if response.status_code == 200:
                return response
            
            # 429错误：速率限制，固定等待15秒
            elif response.status_code == 429:
                # 固定等待15秒
                wait_time = 15
                
                if attempt < max_retries - 1:
                    print(f"  遇到速率限制（429），等待 {wait_time} 秒后重试... (尝试 {attempt + 1}/{max_retries})")
                    print(f"  提示：Semantic Scholar API有速率限制，请耐心等待...")
                    time.sleep(wait_time)
                    continue
                else:
                    print(f"  达到最大重试次数，仍遇到速率限制。")
                    print(f"  建议：请稍后再试，或减少搜索数量，或仅使用Arxiv数据源。")
                    return None
            
            # 其他HTTP错误
            else:
                print(f"  API请求失败，状态码: {response.status_code}")
                if attempt < max_retries - 1:
                    wait_time = min(30, retry_delay * (2 ** attempt))  # 其他错误最多等待30秒
                    print(f"  等待 {wait_time} 秒后重试... (尝试 {attempt + 1}/{max_retries})")
                    time.sleep(wait_time)
                    continue
                return None
                
        except (requests.exceptions.ConnectionError, 
                requests.exceptions.Timeout,
                ConnectionResetError,
                OSError) as e:
            # 连接相关错误，使用更长的等待时间
            wait_time = max(3, retry_delay * (2 ** attempt))
            if attempt < max_retries - 1:
                error_type = type(e).__name__
                print(f"  连接错误 ({error_type})，等待 {wait_time} 秒后重试... (尝试 {attempt + 1}/{max_retries})")
                time.sleep(wait_time)
                continue
            else:
                print(f"  连接失败，已达到最大重试次数。请检查网络连接后重试。")
                return None
                
        except Exception as e:
            # 其他未知错误
            wait_time = min(30, retry_delay * (2 ** attempt))
            if attempt < max_retries - 1:
                print(f"  请求出错: {e}，等待 {wait_time} 秒后重试... (尝试 {attempt + 1}/{max_retries})")
                time.sleep(wait_time)
                continue
            else:
                print(f"  请求失败，已达到最大重试次数: {e}")
                return None
    
    return None

def get_paper_unique_id(paper_info):
    """
    生成论文的唯一标识符，用于去重
    优先级：paperId > title+author > title
    """
    # 如果有paperId，直接使用
    if paper_info.get('paper_id'):
        return f"id:{paper_info['paper_id']}"
    
    # 否则使用title+author的组合
    title = paper_info.get('title', '').lower().strip()
    authors = paper_info.get('authors', [])
    if authors:
        # 取前两个作者的首字母和标题
        author_str = ','.join(sorted([a.lower().strip() for a in authors[:2]]))
        return f"title_author:{hashlib.md5((title + author_str).encode()).hexdigest()}"
    
    # 最后只使用标题
    return f"title:{hashlib.md5(title.encode()).hexdigest()}"

def get_semantic_scholar_paper_details(paper_id):
    """
    通过paperId获取Semantic Scholar论文的详细信息
    包括PDF链接、引用数等
    """
    try:
        url = f"https://api.semanticscholar.org/graph/v1/paper/{paper_id}"
        params = {
            'fields': 'title,authors,year,abstract,url,venue,isOpenAccess,openAccessPdf,externalIds,citationCount,referenceCount'
        }
        response = make_semantic_scholar_request(url, params=params)
        if response:
            return response.json()
    except Exception as e:
        print(f"  获取论文详情时出错: {e}")
    return None

def search_semantic_scholar_papers(keyword, max_results=5, include_details=False, latest_count=6, cited_count=14):
    """
    使用Semantic Scholar API搜索相关论文
    支持分别搜索最新日期和最高引用的论文
    
    处理流程：
    1. 通过关键词搜索论文列表，获取基本信息（标题、作者、摘要、年份、venue等）
    2. 如果include_details=True，通过paperId获取详细信息（PDF链接、引用数等）
    3. 提取并统一格式化论文信息
    4. 注意：Semantic Scholar不提供LaTeX源码，只有摘要和元数据
    
    Args:
        keyword: 搜索关键词
        max_results: 最大结果数量（如果latest_count和cited_count都指定，则忽略此参数）
        include_details: 是否获取详细信息（包括PDF链接、引用数等）
        latest_count: 搜索最新日期的论文数量（默认6篇）
        cited_count: 搜索最高引用的论文数量（默认14篇）
    
    Returns:
        论文信息列表，每个论文包含：标题、作者、发表时间、摘要、venue、paperId、URL等
    """
    print(f"正在从Semantic Scholar搜索关键词: {keyword}...")
    
    papers = []
    all_papers_data = []
    
    try:
        # 如果指定了latest_count和cited_count，分别搜索
        if latest_count > 0:
            print(f"  搜索最新日期的论文（{latest_count}篇）...")
            url = "https://api.semanticscholar.org/graph/v1/paper/search"
            params = {
                'query': keyword,
                'limit': latest_count,
                'fields': 'title,authors,year,abstract,paperId,url,venue,isOpenAccess,citationCount',
                'sort': 'year:desc'  # 按年份降序（最新）
            }
            response = make_semantic_scholar_request(url, params=params)
            if response:
                data = response.json()
                latest_papers_data = data.get('data', [])
                all_papers_data.extend(latest_papers_data)
                print(f"  找到 {len(latest_papers_data)} 篇最新日期的论文")
        
        if cited_count > 0:
            print(f"  搜索最高引用的论文（{cited_count}篇）...")
            url = "https://api.semanticscholar.org/graph/v1/paper/search"
            params = {
                'query': keyword,
                'limit': cited_count,
                'fields': 'title,authors,year,abstract,paperId,url,venue,isOpenAccess,citationCount',
                'sort': 'citationCount:desc'  # 按引用数降序
            }
            response = make_semantic_scholar_request(url, params=params)
            if response:
                data = response.json()
                cited_papers_data = data.get('data', [])
                all_papers_data.extend(cited_papers_data)
                print(f"  找到 {len(cited_papers_data)} 篇最高引用的论文")
        
        # 如果都没有指定，使用原来的逻辑
        if latest_count == 0 and cited_count == 0:
            url = "https://api.semanticscholar.org/graph/v1/paper/search"
            params = {
                'query': keyword,
                'limit': max_results,
                'fields': 'title,authors,year,abstract,paperId,url,venue,isOpenAccess'
            }
            response = make_semantic_scholar_request(url, params=params)
            if response:
                data = response.json()
                all_papers_data = data.get('data', [])
        
        # 处理所有找到的论文
        for paper_data in all_papers_data:
            # 提取作者信息
            authors = []
            if paper_data.get('authors'):
                authors = [author.get('name', '') for author in paper_data['authors']]
            
            # 基础论文信息
            paper_id = paper_data.get('paperId', '')
            paper_info = {
                'title': paper_data.get('title', ''),
                'authors': authors,
                'published': str(paper_data.get('year', '')) if paper_data.get('year') else '',
                'summary': paper_data.get('abstract', ''),
                'entry_id': paper_data.get('url', ''),
                'paper_id': paper_id,
                'venue': paper_data.get('venue', ''),
                'source': 'semantic_scholar',
                'pdf_url': None,
                'latex_content': None,  # Semantic Scholar不提供LaTeX源码
                'citation_count': paper_data.get('citationCount'),  # 如果API返回了引用数，直接使用
                'reference_count': None
            }
            
            # 如果需要详细信息，通过paperId获取
            if include_details and paper_id:
                print(f"  正在获取论文详细信息: {paper_info['title'][:50]}...")
                details = get_semantic_scholar_paper_details(paper_id)
                if details:
                    # 获取PDF链接（如果开放获取）
                    if details.get('isOpenAccess') and details.get('openAccessPdf'):
                        paper_info['pdf_url'] = details.get('openAccessPdf', {}).get('url', '')
                    
                    # 获取引用数（如果之前没有）
                    if paper_info['citation_count'] is None:
                        paper_info['citation_count'] = details.get('citationCount', 0)
                    paper_info['reference_count'] = details.get('referenceCount', 0)
                    
                    # 如果有Arxiv ID，记录以便后续可能获取LaTeX
                    external_ids = details.get('externalIds', {})
                    if external_ids.get('ArXiv'):
                        paper_info['arxiv_id'] = external_ids['ArXiv']
            
            papers.append(paper_info)
            print(f"  找到论文: {paper_info['title']}")
            if paper_info.get('citation_count') is not None:
                print(f"    引用数: {paper_info['citation_count']}, 参考文献数: {paper_info.get('reference_count', 'N/A')}")
            
            # 每找到一篇论文后空闲2秒，以控制访问速度
            time.sleep(2.0)
    except Exception as e:
        print(f"  搜索Semantic Scholar时出错: {e}")
    
    return papers

def search_arxiv_papers(keyword, max_results=5, include_latex=False):
    """
    使用Arxiv API搜索相关论文，按最新日期排序
    """
    print(f"正在搜索关键词: {keyword}...")
    
    papers = []
    try:
        # 使用arxiv库直接搜索，按最新提交日期排序
        search = arxiv.Search(
            query=keyword,
            max_results=max_results,
            sort_by=arxiv.SortCriterion.SubmittedDate,
            sort_order=arxiv.SortOrder.Descending
        )
        
        for paper in search.results():
            try:
                # 处理日期可能为None的情况
                published_date = ''
                if paper.published:
                    published_date = paper.published.strftime('%Y-%m-%d')
                
                paper_info = {
                    'title': paper.title or '',
                    'authors': [author.name for author in paper.authors] if paper.authors else [],
                    'published': published_date,
                    'summary': paper.summary or '',
                    'entry_id': paper.entry_id or '',
                    'paper_id': None,  # Arxiv没有paperId
                    'venue': 'arXiv',
                    'source': 'arxiv',
                    'pdf_url': paper.pdf_url if hasattr(paper, 'pdf_url') else None,
                    'latex_content': None,
                    'citation_count': None
                }
                
                # 如果需要获取LaTeX源码
                if include_latex:
                    latex_content = get_paper_latex_content(paper.entry_id)
                    if latex_content:
                        paper_info['latex_content'] = latex_content
                        print(f"  ✓ 已获取LaTeX源码")
                    else:
                        print(f"  ✗ 无法获取LaTeX源码，将使用摘要")
                
                papers.append(paper_info)
                print(f"  找到论文: {paper_info['title']}")
            except Exception as e:
                print(f"  处理论文时出错: {e}")
                continue
    except Exception as e:
        print(f"搜索Arxiv论文时出错: {e}")
    
    return papers

class OpenAlexClient:
    """
    OpenAlex API客户端，用于搜索论文并获取元数据
    """
    def __init__(self):
        self.base_url = "https://api.openalex.org"
    
    def search_papers(self, query: str, max_results: int = 50, sort: str = "cited_by_count:desc") -> List[Dict]:
        """
        通过OpenAlex搜索论文并获取丰富元数据
        
        Args:
            query: 搜索关键词
            max_results: 最大返回数量
            sort: 排序方式，默认按引用数降序
        
        Returns:
            论文列表
        """
        papers = []
        try:
            url = f"{self.base_url}/works"
            params = {
                'search': query,
                'per-page': min(max_results, 200),  # OpenAlex单次最大200
                'sort': sort
            }
            
            print(f"  正在通过OpenAlex搜索: {query}...")
            response = requests.get(url, params=params, timeout=30)
            response.raise_for_status()
            data = response.json()
            
            for i, work in enumerate(data.get('results', []), 1):
                try:
                    paper_info = self._parse_openalex_paper(work)
                    if paper_info:
                        papers.append(paper_info)
                        print(f"  找到论文 [{i}]: {paper_info['title'][:60]}...")
                        if paper_info.get('citation_count') is not None:
                            print(f"    引用数: {paper_info['citation_count']}")
                        
                        # 每找到一篇论文后空闲2秒，以控制访问速度
                        if i < len(data.get('results', [])):
                            time.sleep(2.0)
                except Exception as e:
                    print(f"  解析论文时出错: {e}")
                    continue
                    
        except Exception as e:
            print(f"  OpenAlex搜索失败: {e}")
        
        return papers
    
    def _parse_openalex_paper(self, work: Dict) -> Optional[Dict]:
        """
        解析OpenAlex返回的论文数据，转换为统一格式
        """
        try:
            # 提取开放获取信息
            oa_info = work.get('open_access', {})
            is_oa = oa_info.get('oa_status', 'closed') != 'closed'
            oa_url = oa_info.get('oa_url')
            
            # 提取全文位置信息
            locations = work.get('locations', [])
            pdf_url = None
            landing_url = None
            
            for location in locations:
                if location.get('pdf_url'):
                    pdf_url = location['pdf_url']
                if location.get('landing_page_url'):
                    landing_url = location['landing_page_url']
            
            # 提取主题信息
            topics = [topic.get('display_name', '') for topic in work.get('topics', [])[:5]]
            
            # 提取作者信息
            authors_list = []
            for authorship in work.get('authorships', []):
                author = authorship.get('author', {})
                if author:
                    authors_list.append(author.get('display_name', ''))
            
            # 提取DOI
            doi = work.get('doi')
            if doi and doi.startswith('https://doi.org/'):
                doi = doi.replace('https://doi.org/', '')
            
            # 提取Arxiv ID（如果有）
            arxiv_id = None
            primary_location = work.get('primary_location', {})
            if primary_location:
                source = primary_location.get('source', {})
                if source and 'arxiv' in source.get('display_name', '').lower():
                    # 尝试从URL提取Arxiv ID
                    landing = primary_location.get('landing_page_url', '')
                    arxiv_match = re.search(r'arxiv\.org/abs/(\d+\.\d+)', landing)
                    if arxiv_match:
                        arxiv_id = arxiv_match.group(1)
            
            # 提取发表信息
            venue = None
            if work.get('primary_location', {}).get('source'):
                venue = work.get('primary_location', {}).get('source', {}).get('display_name', '')
            
            publication_date = work.get('publication_date', '')
            publication_year = work.get('publication_year')
            
            return {
                'title': work.get('title', ''),
                'authors': authors_list[:10],  # 最多10个作者
                'published': publication_date,
                'publication_year': publication_year,
                'summary': work.get('abstract', ''),
                'doi': doi,
                'venue': venue or 'Unknown',
                'source': 'openalex',
                'paper_id': work.get('id', '').replace('https://openalex.org/', '') if work.get('id') else None,
                'entry_id': work.get('id', '') if work.get('id') else None,
                'pdf_url': pdf_url,
                'citation_count': work.get('cited_by_count', 0),
                'reference_count': None,  # OpenAlex不直接提供
                'latex_content': None,
                # OpenAlex特有字段
                'open_access': {
                    'is_oa': is_oa,
                    'oa_status': oa_info.get('oa_status', 'closed'),
                    'oa_url': oa_url
                },
                'full_text_locations': {
                    'pdf_url': pdf_url,
                    'landing_page_url': landing_url
                },
                'topics': topics,
                'arxiv_id': arxiv_id,
                'keywords': [kw.get('display_name', '') for kw in work.get('keywords', [])]
            }
        except Exception as e:
            print(f"  解析OpenAlex论文数据时出错: {e}")
            return None

class FullTextDownloader:
    """
    智能全文下载器，根据论文元数据尝试获取LaTeX/XML全文内容
    """
    def __init__(self):
        self.priority_sources = ['arxiv', 'pubmed', 'acl', 'openreview']
    
    def smart_get_fulltext(self, paper_metadata: Dict) -> Dict:
        """
        根据论文元数据智能获取全文内容
        
        Args:
            paper_metadata: 论文元数据字典
        
        Returns:
            包含content_type, content, source, reason的字典
        """
        content_type = None
        fulltext_content = None
        source = None
        reason = None
        
        # 只有开放获取论文才尝试下载全文
        open_access = paper_metadata.get('open_access', {})
        if isinstance(open_access, dict):
            is_oa = open_access.get('is_oa', False)
        else:
            # 兼容旧格式
            is_oa = paper_metadata.get('source') == 'arxiv'  # Arxiv默认开放获取
        
        if not is_oa and paper_metadata.get('source') != 'arxiv':
            return {
                'content_type': 'metadata_only',
                'content': None,
                'source': None,
                'reason': '非开放获取论文'
            }
        
        # 策略1: 尝试从arXiv获取LaTeX源码
        arxiv_content = self._try_get_arxiv_content(paper_metadata)
        if arxiv_content:
            return {
                'content_type': 'latex',
                'content': arxiv_content,
                'source': 'arxiv',
                'reason': '成功获取LaTeX源码'
            }
        
        # 策略2: 尝试从PubMed Central获取XML
        pmc_content = self._try_get_pmc_content(paper_metadata)
        if pmc_content:
            return {
                'content_type': 'xml',
                'content': pmc_content,
                'source': 'pmc',
                'reason': '成功获取XML全文'
            }
        
        # 策略3: 尝试从ACL Anthology获取LaTeX
        acl_content = self._try_get_acl_content(paper_metadata)
        if acl_content:
            return {
                'content_type': 'latex',
                'content': acl_content,
                'source': 'acl',
                'reason': '成功获取ACL LaTeX源码'
            }
        
        # 策略4: 如果有PDF链接，可以尝试PDF解析（可选，需要pymupdf）
        pdf_url = paper_metadata.get('pdf_url') or \
                  (paper_metadata.get('full_text_locations', {}).get('pdf_url') if isinstance(paper_metadata.get('full_text_locations'), dict) else None)
        if pdf_url and self._is_trusted_pdf_source(pdf_url):
            pdf_text = self._try_parse_pdf(pdf_url)
            if pdf_text:
                return {
                    'content_type': 'pdf_text',
                    'content': pdf_text,
                    'source': 'pdf_parser',
                    'reason': '从PDF解析文本'
                }
        
        # 所有方法都失败，回退到仅使用元数据
        return {
            'content_type': 'metadata_only',
            'content': None,
            'source': None,
            'reason': '无法获取全文内容'
        }
    
    def _try_get_arxiv_content(self, paper_metadata: Dict) -> Optional[str]:
        """尝试从arXiv获取LaTeX源码"""
        try:
            # 检查是否有arxiv_id
            arxiv_id = paper_metadata.get('arxiv_id')
            if not arxiv_id:
                # 尝试从entry_id或DOI提取
                entry_id = paper_metadata.get('entry_id', '')
                if 'arxiv' in entry_id.lower():
                    arxiv_match = re.search(r'arxiv\.org/abs/(\d+\.\d+)', entry_id)
                    if arxiv_match:
                        arxiv_id = arxiv_match.group(1)
                
                if not arxiv_id:
                    doi = paper_metadata.get('doi', '')
                    if 'arxiv' in doi.lower():
                        arxiv_match = re.search(r'arxiv\.(\d+\.\d+)', doi)
                        if arxiv_match:
                            arxiv_id = arxiv_match.group(1)
            
            if arxiv_id:
                # 使用现有的get_paper_latex_content函数
                latex_content = get_paper_latex_content(f"http://arxiv.org/abs/{arxiv_id}")
                if latex_content:
                    return latex_content
        except Exception as e:
            print(f"    获取arXiv内容失败: {e}")
        return None
    
    def _try_get_pmc_content(self, paper_metadata: Dict) -> Optional[str]:
        """尝试从PubMed Central获取XML全文"""
        try:
            doi = paper_metadata.get('doi')
            if not doi:
                return None
            
            # 通过PMC API获取XML内容
            # 首先需要通过DOI获取PMC ID
            pmc_id = self._get_pmc_id_by_doi(doi)
            if pmc_id:
                # 下载PMC全文XML
                pmc_url = f"https://www.ncbi.nlm.nih.gov/pmc/oai/oai.cgi?verb=GetRecord&identifier=oai:pubmedcentral.nih.gov:{pmc_id}&metadataPrefix=pmc"
                response = requests.get(pmc_url, timeout=30)
                if response.status_code == 200:
                    # 这里需要解析XML，简化处理
                    return response.text[:50000]  # 限制长度
        except Exception as e:
            print(f"    获取PMC内容失败: {e}")
        return None
    
    def _get_pmc_id_by_doi(self, doi: str) -> Optional[str]:
        """通过DOI获取PMC ID"""
        try:
            url = f"https://www.ncbi.nlm.nih.gov/pmc/utils/idconv/v1.0/?ids={doi}&format=json"
            response = requests.get(url, timeout=30)
            if response.status_code == 200:
                data = response.json()
                records = data.get('records', [])
                if records and records[0].get('pmcid'):
                    return records[0]['pmcid'].replace('PMC', '')
        except Exception as e:
            pass
        return None
    
    def _try_get_acl_content(self, paper_metadata: Dict) -> Optional[str]:
        """尝试从ACL Anthology获取LaTeX源码"""
        try:
            # 检查是否是ACL相关会议
            title = paper_metadata.get('title', '').lower()
            venue = paper_metadata.get('venue', '').lower()
            topics = paper_metadata.get('topics', [])
            
            acl_keywords = ['acl', 'emnlp', 'naacl', 'eacl', 'computational linguistics', 'natural language processing']
            if any(keyword in title for keyword in acl_keywords) or \
               any(keyword in venue for keyword in acl_keywords) or \
               any(topic.lower() in acl_keywords for topic in topics if isinstance(topic, str)):
                # ACL Anthology通常可以通过DOI或标题搜索
                # 这里简化处理，实际需要访问ACL Anthology API
                pass
        except Exception as e:
            print(f"    获取ACL内容失败: {e}")
        return None
    
    def _is_trusted_pdf_source(self, pdf_url: str) -> bool:
        """判断是否是可信的PDF源"""
        trusted_domains = ['arxiv.org', 'pubmedcentral.nih.gov', 'aclweb.org', 'openreview.net']
        return any(domain in pdf_url.lower() for domain in trusted_domains)
    
    def _try_parse_pdf(self, pdf_url: str) -> Optional[str]:
        """尝试从PDF URL解析文本内容（需要pymupdf库）"""
        try:
            try:
                import pymupdf  # PyMuPDF
            except ImportError:
                # 如果没有安装pymupdf，跳过PDF解析
                return None
            
            import io
            response = requests.get(pdf_url, timeout=30)
            if response.status_code == 200:
                pdf_data = io.BytesIO(response.content)
                doc = pymupdf.open(stream=pdf_data, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
                return text if len(text) > 100 else None  # 确保有足够内容
        except Exception as e:
            print(f"    PDF解析失败: {e}")
        return None

def search_openalex_papers(keyword, max_results=50, include_details=False, latest_count=6, cited_count=14):
    """
    使用OpenAlex API搜索相关论文
    支持分别搜索最新日期和最高引用的论文
    
    Args:
        keyword: 搜索关键词
        max_results: 最大返回数量（如果latest_count和cited_count都指定，则忽略此参数）
        include_details: 是否获取详细信息（OpenAlex默认包含详细信息）
        latest_count: 搜索最新日期的论文数量（默认6篇）
        cited_count: 搜索最高引用的论文数量（默认14篇）
    
    Returns:
        论文列表
    """
    client = OpenAlexClient()
    papers = []
    
    # 如果指定了latest_count和cited_count，分别搜索
    if latest_count > 0:
        print(f"  搜索最新日期的论文（{latest_count}篇）...")
        latest_papers = client.search_papers(keyword, max_results=latest_count, sort="publication_date:desc")
        papers.extend(latest_papers)
        print(f"  找到 {len(latest_papers)} 篇最新日期的论文")
    
    if cited_count > 0:
        print(f"  搜索最高引用的论文（{cited_count}篇）...")
        cited_papers = client.search_papers(keyword, max_results=cited_count, sort="cited_by_count:desc")
        papers.extend(cited_papers)
        print(f"  找到 {len(cited_papers)} 篇最高引用的论文")
    
    # 如果都没有指定，使用原来的逻辑
    if latest_count == 0 and cited_count == 0:
        papers = client.search_papers(keyword, max_results=max_results)
    
    return papers

def merge_and_deduplicate_papers(arxiv_papers, semantic_papers, openalex_papers=None):
    """
    合并Arxiv、Semantic Scholar和OpenAlex的搜索结果，并去除重复论文
    
    Args:
        arxiv_papers: Arxiv论文列表
        semantic_papers: Semantic Scholar论文列表
        openalex_papers: OpenAlex论文列表（可选）
    
    Returns:
        合并去重后的论文列表
    """
    all_papers = []
    seen_papers: Set[str] = set()
    duplicate_count = 0
    
    # 先添加Arxiv的论文
    for paper in arxiv_papers:
        unique_id = get_paper_unique_id(paper)
        if unique_id not in seen_papers:
            seen_papers.add(unique_id)
            all_papers.append(paper)
        else:
            duplicate_count += 1
            print(f"  去重: 跳过重复论文 - {paper['title']}")
    
    # 再添加Semantic Scholar的论文（去重）
    for paper in semantic_papers:
        unique_id = get_paper_unique_id(paper)
        if unique_id not in seen_papers:
            seen_papers.add(unique_id)
            all_papers.append(paper)
        else:
            duplicate_count += 1
            print(f"  去重: 跳过重复论文 - {paper['title']}")
    
    # 最后添加OpenAlex的论文（去重）
    if openalex_papers:
        for paper in openalex_papers:
            unique_id = get_paper_unique_id(paper)
            if unique_id not in seen_papers:
                seen_papers.add(unique_id)
                all_papers.append(paper)
            else:
                duplicate_count += 1
                print(f"  去重: 跳过重复论文 - {paper['title']}")
    
    if duplicate_count > 0:
        print(f"\n  去重统计: 共发现 {duplicate_count} 篇重复论文，已自动去除")
    
    return all_papers

def parse_latex_to_text(latex_content):
    """
    将LaTeX源码解析为可读的文本内容
    提取章节标题、正文、公式说明等，去除LaTeX命令
    """
    if not latex_content:
        return None
    
    # 提取主要.tex文件的内容（通常是主文件）
    main_content = latex_content
    
    # 1. 提取标题
    title_match = re.search(r'\\title\{(.*?)\}', main_content, re.DOTALL)
    title = title_match.group(1).strip() if title_match else ""
    # 清理标题中的LaTeX命令
    title = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', title)
    title = re.sub(r'\\[a-zA-Z]+', '', title)
    
    # 2. 提取摘要
    abstract_match = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', main_content, re.DOTALL)
    abstract = abstract_match.group(1).strip() if abstract_match else ""
    
    # 3. 提取章节结构
    sections = []
    # 查找所有section、subsection等
    section_pattern = r'\\(sub)?section\*?\{([^}]+)\}'
    for match in re.finditer(section_pattern, main_content):
        section_type = "小节" if match.group(1) else "章节"
        section_title = match.group(2).strip()
        sections.append(f"{section_type}: {section_title}")
    
    # 4. 提取正文内容（去除LaTeX命令）
    # 移除注释
    text = re.sub(r'%.*?\n', '\n', main_content)
    # 移除LaTeX命令（保留参数内容）
    text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    # 移除环境标签但保留内容
    text = re.sub(r'\\begin\{[^}]+\}', '', text)
    text = re.sub(r'\\end\{[^}]+\}', '', text)
    # 移除标签和引用
    text = re.sub(r'\\label\{[^}]+\}', '', text)
    text = re.sub(r'\\cite\{[^}]+\}', '[引用]', text)
    text = re.sub(r'\\ref\{[^}]+\}', '[引用]', text)
    # 移除公式环境（保留公式说明）
    text = re.sub(r'\$.*?\$', '[公式]', text)
    text = re.sub(r'\\begin\{equation\}.*?\\end\{equation\}', '[公式]', text, flags=re.DOTALL)
    # 清理多余空白
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    
    # 提取前10000字符作为主要文本内容
    main_text = text[:10000].strip()
    
    # 5. 构建结构化输出
    parsed_content = ""
    if title:
        parsed_content += f"标题: {title}\n\n"
    if abstract:
        parsed_content += f"摘要: {abstract}\n\n"
    if sections:
        parsed_content += f"论文结构:\n" + "\n".join(sections[:10]) + "\n\n"
    if main_text:
        parsed_content += f"主要文本内容:\n{main_text}\n"
    
    return parsed_content if parsed_content.strip() else None

def extract_arxiv_id(entry_id):
    """
    从entry_id中提取Arxiv ID，处理可能包含版本号的情况
    例如：http://arxiv.org/abs/1234.5678v1 -> 1234.5678
    """
    if not entry_id:
        return None
    # 提取最后一个路径段
    arxiv_id = entry_id.split('/')[-1]
    # 移除版本号（如v1, v2等）
    if 'v' in arxiv_id and arxiv_id[-1].isdigit():
        # 找到最后一个v的位置
        last_v = arxiv_id.rfind('v')
        if last_v > 0:
            arxiv_id = arxiv_id[:last_v]
    return arxiv_id

def get_citation_count_for_arxiv_paper(paper_info):
    """
    为Arxiv论文查找引用次数（通过Semantic Scholar API）
    使用Arxiv ID或标题+作者进行搜索
    """
    try:
        # 首先尝试使用Arxiv ID查找
        arxiv_id = None
        if paper_info.get('entry_id'):
            # 从entry_id提取Arxiv ID（格式：http://arxiv.org/abs/1234.5678）
            arxiv_id = extract_arxiv_id(paper_info['entry_id'])
        
        # 方法1: 如果有Arxiv ID，直接通过Arxiv ID查找
        if arxiv_id:
            try:
                # Semantic Scholar支持通过Arxiv ID查找
                url = f"https://api.semanticscholar.org/graph/v1/paper/arXiv:{arxiv_id}"
                params = {
                    'fields': 'citationCount,referenceCount'
                }
                response = make_semantic_scholar_request(url, params=params)
                if response:
                    data = response.json()
                    citation_count = data.get('citationCount', 0)
                    if citation_count is not None:
                        return citation_count
            except Exception as e:
                # 静默失败，继续尝试其他方法
                pass
        
        # 方法2: 通过标题搜索（如果Arxiv ID方法失败）
        title = paper_info.get('title', '')
        if title:
            try:
                url = "https://api.semanticscholar.org/graph/v1/paper/search"
                params = {
                    'query': title,
                    'limit': 5,
                    'fields': 'title,authors,citationCount'
                }
                response = make_semantic_scholar_request(url, params=params)
                if response:
                    data = response.json()
                    papers_data = data.get('data', [])
                    # 尝试匹配标题
                    for match in papers_data:
                        if match.get('title', '').lower() == title.lower():
                            citation_count = match.get('citationCount', 0)
                            if citation_count is not None:
                                return citation_count
            except Exception as e:
                # 静默失败，继续返回None
                pass
    except Exception as e:
        print(f"  获取引用次数时出错: {e}")
    
    return None

def enrich_papers_with_citations(papers):
    """
    为论文补充引用次数信息
    仅对 OpenAlex 与 Semantic Scholar 来源的论文进行处理；跳过 Arxiv。
    """
    print("\n正在获取论文引用次数信息（仅OpenAlex与Semantic Scholar，跳过Arxiv）...")
    processed = 0
    for i, paper in enumerate(papers):
        source = (paper.get('source') or '').lower()
        # 跳过 Arxiv 来源
        if source == 'arxiv':
            # 若无引用次数字段，统一初始化为0，避免后续使用时报错
            if paper.get('citation_count') is None:
                paper['citation_count'] = 0
            continue
        # 仅处理 OpenAlex 与 Semantic Scholar
        if source in ('openalex', 'semantic_scholar', 'semanticscholar'):
            # 避免不必要的外部请求：若已携带引用次数则标准化类型后跳过
            if paper.get('citation_count') is not None:
                try:
                    paper['citation_count'] = int(paper['citation_count']) or 0
                except Exception:
                    paper['citation_count'] = 0
                continue
            # Semantic Scholar 可能需要补充细节
            if source in ('semantic_scholar', 'semanticscholar') and paper.get('paper_id'):
                details = get_semantic_scholar_paper_details(paper['paper_id'])
                if details:
                    citation_count = details.get('citationCount', 0)
                    paper['citation_count'] = citation_count if citation_count is not None else 0
                    reference_count = details.get('referenceCount', 0)
                    paper['reference_count'] = reference_count if reference_count is not None else 0
                    print(f"  ✓ {paper['title'][:50]}... 引用数: {paper['citation_count']}")
                else:
                    paper['citation_count'] = 0
            else:
                # OpenAlex 通常已包含引用数；若无则设为0
                paper['citation_count'] = int(paper.get('citation_count') or 0)
            processed += 1
        else:
            # 其他来源统一不查询，规范化字段
            if paper.get('citation_count') is None:
                paper['citation_count'] = 0
    if processed == 0:
        print("  - 未发现需要获取引用次数的论文（OpenAlex/Semantic Scholar）。")
    return papers

def filter_papers_by_criteria(papers, top_by_date=5, top_by_citations=5, min_citations=10):
    """
    根据最新日期和引用次数筛选论文
    
    Args:
        papers: 论文列表
        top_by_date: 按最新日期选择前N篇
        top_by_citations: 按引用次数选择前N篇（即使日期不新）
        min_citations: 高引用论文的最低引用数阈值
    
    Returns:
        (latest_papers, high_citation_papers, all_selected)
        latest_papers: 按日期筛选的论文
        high_citation_papers: 按引用次数筛选的论文
        all_selected: 合并去重后的所有选中论文
    """
    # 解析日期
    def parse_date(date_str):
        if not date_str:
            return datetime(1900, 1, 1)
        try:
            # 尝试多种日期格式
            if isinstance(date_str, str):
                # YYYY-MM-DD格式
                if '-' in date_str and len(date_str) >= 10:
                    return datetime.strptime(date_str[:10], '%Y-%m-%d')
                # YYYY格式
                elif date_str.isdigit() and len(date_str) == 4:
                    return datetime(int(date_str), 1, 1)
            return datetime(1900, 1, 1)
        except (ValueError, TypeError) as e:
            # 日期解析失败，返回默认日期
            return datetime(1900, 1, 1)
    
    # 为每篇论文计算解析后的日期和引用次数（不修改原数据）
    papers_with_metadata = []
    for paper in papers:
        paper_copy = paper.copy()  # 创建副本以避免污染原数据
        paper_copy['_parsed_date'] = parse_date(paper.get('published', ''))
        paper_copy['_citation_count'] = int(paper.get('citation_count', 0) or 0)
        papers_with_metadata.append(paper_copy)
    
    # 1. 按最新日期筛选
    sorted_by_date = sorted(papers_with_metadata, key=lambda x: x['_parsed_date'], reverse=True)
    latest_papers = [p.copy() for p in sorted_by_date[:top_by_date]]
    # 移除临时字段
    for p in latest_papers:
        p.pop('_parsed_date', None)
        p.pop('_citation_count', None)
    
    # 2. 按引用次数筛选（排除已在最新日期列表中的）
    latest_ids = {get_paper_unique_id(p) for p in latest_papers}
    remaining_papers = [p for p in papers_with_metadata if get_paper_unique_id(p) not in latest_ids]
    
    # 筛选高引用论文（引用数 >= min_citations）
    high_citation_papers = [p for p in remaining_papers if p['_citation_count'] >= min_citations]
    high_citation_papers = sorted(high_citation_papers, key=lambda x: x['_citation_count'], reverse=True)[:top_by_citations]
    # 移除临时字段
    for p in high_citation_papers:
        p.pop('_parsed_date', None)
        p.pop('_citation_count', None)
    
    # 3. 合并去重
    all_selected = []
    seen_ids = set()
    for paper in latest_papers + high_citation_papers:
        unique_id = get_paper_unique_id(paper)
        if unique_id not in seen_ids:
            seen_ids.add(unique_id)
            all_selected.append(paper)
    
    return latest_papers, high_citation_papers, all_selected

def display_papers_for_selection(papers, category_name=""):
    """
    显示论文列表供用户选择
    """
    if not papers:
        return []
    
    print(f"\n{'='*80}")
    if category_name:
        print(f"{category_name} ({len(papers)}篇)")
    print(f"{'='*80}")
    
    for i, paper in enumerate(papers, 1):
        print(f"\n[{i}] {paper.get('title', '未知标题')}")
        authors = paper.get('authors') or []
        print(f"    作者: {', '.join(authors[:3]) if authors else '未知'}")
        print(f"    发表时间: {paper.get('published', '未知')}")
        if paper.get('citation_count') is not None:
            print(f"    引用数: {paper['citation_count']}")
        if paper.get('venue'):
            print(f"    来源: {paper['venue']} ({paper.get('source', 'unknown')})")
        summary = paper.get('summary')
        if summary and isinstance(summary, str):
            summary_display = summary[:150] + "..." if len(summary) > 150 else summary
            print(f"    摘要: {summary_display}")
    
    print(f"\n{'='*80}")
    print(f"请选择要包含在综述中的论文（输入编号，多个用逗号分隔，如：1,3,5，或输入 'all' 选择全部）:")
    choice = input().strip()
    
    if not choice:
        print("未输入任何内容，未选择任何论文")
        return []
    
    if choice.lower() == 'all':
        return papers
    else:
        try:
            indices = [int(x.strip()) - 1 for x in choice.split(',') if x.strip().isdigit()]
            if not indices:
                print("输入格式错误，未选择任何论文")
                return []
            selected = [papers[i] for i in indices if 0 <= i < len(papers)]
            if len(selected) < len(indices):
                print(f"警告：部分编号无效（有效范围：1-{len(papers)}）")
            return selected
        except (ValueError, IndexError) as e:
            print(f"输入格式错误，未选择任何论文: {e}")
            return []

def get_paper_latex_content(paper_id):
    """
    获取论文的LaTeX源码内容并解析为可读文本
    """
    try:
        # 从entry_id中提取arxiv ID（格式：http://arxiv.org/abs/1234.5678）
        arxiv_id = extract_arxiv_id(paper_id)
        if not arxiv_id:
            print(f"  无法从entry_id提取Arxiv ID: {paper_id}")
            return None
        
        # Arxiv的LaTeX源码URL
        latex_url = f"https://arxiv.org/e-print/{arxiv_id}"
        
        print(f"正在下载LaTeX源码: {arxiv_id}...")
        
        # 下载LaTeX源码（通常是tar.gz格式）
        response = requests.get(latex_url, timeout=30)
        if response.status_code == 200:
            # 尝试解压并提取.tex文件内容
            try:
                tar_file = tarfile.open(fileobj=io.BytesIO(response.content))
                all_tex_content = ""
                main_tex_file = None
                
                # 首先查找主文件（通常是包含\documentclass的文件）
                for member in tar_file.getmembers():
                    if member.name.endswith('.tex') and member.isfile():
                        try:
                            file_content = tar_file.extractfile(member).read().decode('utf-8', errors='ignore')
                            if '\\documentclass' in file_content:
                                main_tex_file = file_content
                                print(f"  找到主文件: {member.name}")
                            all_tex_content += file_content + "\n\n"
                        except Exception as e:
                            print(f"  读取文件 {member.name} 时出错: {e}")
                
                tar_file.close()
                
                # 优先使用主文件，否则使用所有内容
                latex_source = main_tex_file if main_tex_file else all_tex_content
                
                if latex_source:
                    # 解析LaTeX为可读文本
                    print(f"  正在解析LaTeX内容...")
                    parsed_text = parse_latex_to_text(latex_source)
                    if parsed_text:
                        return parsed_text
                    else:
                        # 如果解析失败，返回原始内容的前20000字符
                        return latex_source[:20000]
                else:
                    return None
            except Exception as e:
                print(f"解压LaTeX源码时出错: {e}")
                return None
        else:
            print(f"下载LaTeX源码失败，状态码: {response.status_code}")
            return None
    except Exception as e:
        print(f"获取LaTeX源码时出错: {e}")
        return None

def select_cited_papers_with_llm(papers, review, target_count, keyword):
    """
    使用大模型根据生成的综述选择要引用的文献
    
    Args:
        papers: 所有论文列表
        review: 生成的综述内容
        target_count: 目标选择的文献数量
        keyword: 搜索关键词
    
    Returns:
        选中的论文列表
    """
    if not papers:
        return []
    
    if len(papers) <= target_count:
        print(f"论文数量({len(papers)})已小于等于目标数量({target_count})，使用所有论文")
        return papers
    
    print(f"\n正在使用大模型选择引用文献（从{len(papers)}篇中选择{target_count}篇）...")
    
    try:
        # 构建论文信息摘要
        papers_content = ""
        for i, paper in enumerate(papers, 1):
            title = paper.get('title') or '未知标题'
            authors = paper.get('authors') or []
            published = paper.get('published') or '未知'
            venue = paper.get('venue') or ''
            source = paper.get('source') or 'unknown'
            citation_count = paper.get('citation_count') or 0
            summary = paper.get('summary') or '无摘要'
            
            papers_content += f"\n论文编号 {i}:\n"
            papers_content += f"标题: {title}\n"
            papers_content += f"作者: {', '.join(authors) if authors else '未知'}\n"
            papers_content += f"发表时间: {published}\n"
            if venue:
                papers_content += f"来源: {venue} ({source})\n"
            papers_content += f"引用数: {citation_count}\n"
            # 只包含摘要的前200字符，避免prompt过长
            if summary and isinstance(summary, str):
                summary_short = summary[:200] + "..." if len(summary) > 200 else summary
            else:
                summary_short = '无摘要'
            papers_content += f"摘要: {summary_short}\n"
            papers_content += "-" * 60 + "\n"
        
        # 构建提示词
        system_prompt = """你是一位专业的学术研究助手，擅长根据综述内容选择最相关的引用文献。

选择标准：
1. 优先选择在综述中被提及或引用的论文
2. 优先选择与综述主题高度相关的论文
3. 优先选择日期较新或引用次数较高的论文
4. 确保选择的文献能够支撑综述的主要观点和技术总结
5. 综合考虑论文的相关性、重要性和代表性

请仔细分析综述内容和每篇论文的信息，选择最合适的引用文献。"""
        
        user_prompt = f"""以下是一篇关于"{keyword}"的学术综述：

{review[:3000]}...

（综述内容已截断，请基于上述内容选择相关文献）

请从以下{len(papers)}篇论文中，选择最相关的{target_count}篇作为引用文献。

选择标准：
1. 优先选择在综述中被提及或引用的论文
2. 优先选择与综述主题高度相关的论文
3. 优先选择日期较新或引用次数较高的论文
4. 确保选择的文献能够支撑综述的主要观点和技术总结

请严格按照以下格式输出，只输出论文编号，用逗号分隔，不要有其他内容：
编号1,编号2,编号3,...

例如：1,3,5,7,9

论文列表：
{papers_content}

请输出选择的{target_count}篇论文的编号（用逗号分隔）："""
        
        # 调用大模型
        response = client.chat.completions.create(
            model='qwen-max',
            messages=[
                {
                    'role': 'system',
                    'content': system_prompt
                },
                {
                    'role': 'user',
                    'content': user_prompt
                }
            ],
            stream=False,
            temperature=0.3  # 使用较低的温度以获得更稳定的选择结果
        )
        
        if response and response.choices and len(response.choices) > 0:
            result_text = response.choices[0].message.content.strip()
            
            # 解析返回的编号
            try:
                # 提取数字编号
                numbers = re.findall(r'\d+', result_text)
                selected_indices = [int(n) - 1 for n in numbers]  # 转换为0-based索引
                
                # 验证索引范围
                valid_indices = [idx for idx in selected_indices if 0 <= idx < len(papers)]
                
                # 去重
                unique_indices = list(dict.fromkeys(valid_indices))  # 保持顺序的去重
                
                # 限制数量
                if len(unique_indices) > target_count:
                    unique_indices = unique_indices[:target_count]
                elif len(unique_indices) < target_count and len(unique_indices) < len(papers):
                    # 如果选择的论文数量不足，补充一些高引用或新日期的论文
                    remaining_indices = [i for i in range(len(papers)) if i not in unique_indices]
                    # 按引用数和日期排序
                    remaining_papers = [(i, papers[i]) for i in remaining_indices]
                    remaining_papers.sort(key=lambda x: (
                        x[1].get('citation_count') or 0,
                        x[1].get('published', '') or ''
                    ), reverse=True)
                    # 补充到目标数量
                    needed = target_count - len(unique_indices)
                    for i, _ in remaining_papers[:needed]:
                        unique_indices.append(i)
                
                selected_papers = [papers[i] for i in unique_indices]
                print(f"成功选择了 {len(selected_papers)} 篇文献")
                return selected_papers
            except Exception as e:
                print(f"解析大模型返回结果时出错: {e}")
                print(f"返回的文本: {result_text}")
                # 使用备用方案：按引用数和日期排序
                return select_cited_papers_fallback(papers, target_count)
        else:
            print("大模型返回格式异常，使用备用选择方案")
            return select_cited_papers_fallback(papers, target_count)
            
    except Exception as e:
        error_msg = f"使用大模型选择引用文献时出错: {e}"
        print(error_msg)
        print("  使用备用选择方案：按日期和引用次数排序")
        return select_cited_papers_fallback(papers, target_count)

def select_cited_papers_fallback(papers, target_count):
    """
    备用选择方案：按日期和引用次数综合排序
    """
    def parse_date(date_str):
        if not date_str:
            return datetime(1900, 1, 1)
        try:
            if isinstance(date_str, str):
                if '-' in date_str and len(date_str) >= 10:
                    return datetime.strptime(date_str[:10], '%Y-%m-%d')
                elif date_str.isdigit() and len(date_str) == 4:
                    return datetime(int(date_str), 1, 1)
            return datetime(1900, 1, 1)
        except (ValueError, TypeError):
            return datetime(1900, 1, 1)
    
    # 为每篇论文计算综合得分：日期得分 + 引用得分
    papers_with_score = []
    for paper in papers:
        paper_copy = paper.copy()
        # 日期得分：越新得分越高（归一化到0-100）
        date = parse_date(paper.get('published', ''))
        if date.year >= 2020:
            date_score = 100 - (2024 - date.year) * 10  # 2024年得100分，每年减10分
        else:
            date_score = max(0, 50 - (2020 - date.year) * 5)  # 2020年之前按年份递减
        
        # 引用得分：引用数越高得分越高（归一化到0-100，使用对数缩放）
        citation_count = int(paper.get('citation_count', 0) or 0)
        if citation_count > 0:
            citation_score = min(100, 20 * (1 + (citation_count ** 0.5)))  # 对数缩放
        else:
            citation_score = 0
        
        # 综合得分：日期权重40%，引用权重60%
        total_score = date_score * 0.4 + citation_score * 0.6
        paper_copy['_score'] = total_score
        papers_with_score.append(paper_copy)
    
    # 按得分排序
    sorted_papers = sorted(papers_with_score, key=lambda x: x['_score'], reverse=True)
    
    # 移除临时字段
    selected_papers = []
    for paper in sorted_papers[:target_count]:
        paper.pop('_score', None)
        selected_papers.append(paper)
    
    return selected_papers

def filter_papers_with_llm(papers, target_count, keyword):
    """
    使用大模型筛选论文，筛选标准是日期较新或引用较高
    
    Args:
        papers: 论文列表
        target_count: 目标筛选出的论文数量
        keyword: 搜索关键词
    
    Returns:
        筛选后的论文列表
    """
    if not papers:
        return []
    
    if len(papers) <= target_count:
        print(f"论文数量({len(papers)})已小于等于目标数量({target_count})，无需筛选")
        return papers
    
    print(f"\n正在使用大模型筛选论文（从{len(papers)}篇中筛选出{target_count}篇）...")
    
    try:
        # 构建论文信息摘要
        papers_content = ""
        for i, paper in enumerate(papers, 1):
            title = paper.get('title') or '未知标题'
            authors = paper.get('authors') or []
            published = paper.get('published') or '未知'
            venue = paper.get('venue') or ''
            source = paper.get('source') or 'unknown'
            citation_count = paper.get('citation_count') or 0
            summary = paper.get('summary') or '无摘要'
            
            papers_content += f"\n论文编号 {i}:\n"
            papers_content += f"标题: {title}\n"
            papers_content += f"作者: {', '.join(authors) if authors else '未知'}\n"
            papers_content += f"发表时间: {published}\n"
            if venue:
                papers_content += f"来源: {venue} ({source})\n"
            papers_content += f"引用数: {citation_count}\n"
            # 只包含摘要的前200字符，避免prompt过长
            if summary and isinstance(summary, str):
                summary_short = summary[:200] + "..." if len(summary) > 200 else summary
            else:
                summary_short = '无摘要'
            papers_content += f"摘要: {summary_short}\n"
            papers_content += "-" * 60 + "\n"
        
        # 构建提示词
        system_prompt = """你是一位专业的学术研究助手，擅长根据论文的日期和引用次数筛选高质量的学术论文。

筛选标准：
1. 优先选择日期较新的论文（最近发表的论文通常代表最新研究进展）
2. 优先选择引用次数较高的论文（高引用通常表示论文质量高、影响力大）
3. 综合考虑日期和引用次数，选择最有价值和代表性的论文

请仔细分析每篇论文的发表时间和引用次数，选择最合适的论文。"""
        
        user_prompt = f"""请从以下关于"{keyword}"的{len(papers)}篇论文中，筛选出最合适的{target_count}篇论文。

筛选标准：优先选择日期较新或引用次数较高的论文。请综合考虑这两个因素，选择最有价值和代表性的论文。

请严格按照以下格式输出，只输出论文编号，用逗号分隔，不要有其他内容：
编号1,编号2,编号3,...

例如：1,3,5,7,9

论文列表：
{papers_content}

请输出筛选出的{target_count}篇论文的编号（用逗号分隔）："""
        
        # 调用大模型
        response = client.chat.completions.create(
            model='qwen-max',
            messages=[
                {
                    'role': 'system',
                    'content': system_prompt
                },
                {
                    'role': 'user',
                    'content': user_prompt
                }
            ],
            stream=False,
            temperature=0.3  # 使用较低的温度以获得更稳定的筛选结果
        )
        
        if response and response.choices and len(response.choices) > 0:
            result_text = response.choices[0].message.content.strip()
            
            # 解析返回的编号
            try:
                # 提取数字编号
                numbers = re.findall(r'\d+', result_text)
                selected_indices = [int(n) - 1 for n in numbers]  # 转换为0-based索引
                
                # 验证索引范围
                valid_indices = [idx for idx in selected_indices if 0 <= idx < len(papers)]
                
                if len(valid_indices) != target_count:
                    print(f"  警告：大模型返回了{len(valid_indices)}个编号，但目标是{target_count}个")
                    # 如果返回的数量不足，补充一些
                    if len(valid_indices) < target_count:
                        remaining = [i for i in range(len(papers)) if i not in valid_indices]
                        needed = target_count - len(valid_indices)
                        valid_indices.extend(remaining[:needed])
                    # 如果返回的数量过多，只取前target_count个
                    elif len(valid_indices) > target_count:
                        valid_indices = valid_indices[:target_count]
                
                selected_papers = [papers[idx] for idx in valid_indices]
                print(f"  筛选完成：从{len(papers)}篇中筛选出{len(selected_papers)}篇论文")
                return selected_papers
                
            except Exception as e:
                print(f"  解析大模型返回结果时出错: {e}")
                print(f"  返回的文本: {result_text}")
                # 如果解析失败，使用备用方案：按日期和引用次数排序后取前N篇
                print("  使用备用筛选方案：按日期和引用次数排序")
                return filter_papers_fallback(papers, target_count)
        else:
            print("  大模型返回格式异常，使用备用筛选方案")
            return filter_papers_fallback(papers, target_count)
            
    except Exception as e:
        error_msg = f"使用大模型筛选论文时出错: {e}"
        print(error_msg)
        print("  使用备用筛选方案：按日期和引用次数排序")
        return filter_papers_fallback(papers, target_count)

def filter_papers_fallback(papers, target_count):
    """
    备用筛选方案：按日期和引用次数综合排序
    """
    def parse_date(date_str):
        if not date_str:
            return datetime(1900, 1, 1)
        try:
            if isinstance(date_str, str):
                if '-' in date_str and len(date_str) >= 10:
                    return datetime.strptime(date_str[:10], '%Y-%m-%d')
                elif date_str.isdigit() and len(date_str) == 4:
                    return datetime(int(date_str), 1, 1)
            return datetime(1900, 1, 1)
        except (ValueError, TypeError):
            return datetime(1900, 1, 1)
    
    # 为每篇论文计算综合得分：日期得分 + 引用得分
    papers_with_score = []
    for paper in papers:
        paper_copy = paper.copy()
        # 日期得分：越新得分越高（归一化到0-100）
        date = parse_date(paper.get('published', ''))
        if date.year >= 2020:
            date_score = 100 - (2024 - date.year) * 10  # 2024年得100分，每年减10分
        else:
            date_score = max(0, 50 - (2020 - date.year) * 5)  # 2020年之前按年份递减
        
        # 引用得分：引用数越高得分越高（归一化到0-100，使用对数缩放）
        citation_count = int(paper.get('citation_count', 0) or 0)
        if citation_count > 0:
            citation_score = min(100, 20 * (1 + (citation_count ** 0.5)))  # 对数缩放
        else:
            citation_score = 0
        
        # 综合得分：日期权重40%，引用权重60%
        total_score = date_score * 0.4 + citation_score * 0.6
        paper_copy['_score'] = total_score
        papers_with_score.append(paper_copy)
    
    # 按得分排序
    sorted_papers = sorted(papers_with_score, key=lambda x: x['_score'], reverse=True)
    
    # 移除临时字段
    selected_papers = []
    for paper in sorted_papers[:target_count]:
        paper.pop('_score', None)
        selected_papers.append(paper)
    
    return selected_papers

def generate_review(papers, keyword, final_citation_count=30):
    """
    使用大模型生成综述，根据内容可用性智能调整综述深度
    生成综述后，由大模型选择要引用的文献
    
    Args:
        papers: 所有论文列表（用于生成综述）
        keyword: 关键词
        final_citation_count: 最终要引用的文献数量（默认30篇）
    
    Returns:
        (review, cited_papers): 综述内容和选中的引用文献列表
    """
    if not papers:
        return ("未提供任何论文，无法生成综述。", [])
    
    print(f"\n正在使用大模型生成综述...")
    
    try:
        # 分类论文：有全文内容和仅元数据
        papers_with_fulltext = []
        papers_metadata_only = []
        
        for paper in papers:
            # 检查是否有全文内容
            full_text = paper.get('full_text', {})
            if isinstance(full_text, dict) and full_text.get('content_type') != 'metadata_only' and full_text.get('content'):
                papers_with_fulltext.append(paper)
            elif paper.get('latex_content'):
                # 兼容旧格式：直接有latex_content的也算有全文
                papers_with_fulltext.append(paper)
            else:
                papers_metadata_only.append(paper)
        
        print(f"  找到 {len(papers_with_fulltext)} 篇可获取全文的论文")
        print(f"  找到 {len(papers_metadata_only)} 篇仅元数据的论文")
        
        # 根据内容可用性选择综述策略
        if len(papers_with_fulltext) >= 10:
            review_strategy = "deep"  # 深度综述
            print("  使用深度综述策略（基于全文内容）")
        elif len(papers_with_fulltext) >= 5:
            review_strategy = "medium"  # 中等深度综述
            print("  使用中等深度综述策略（混合全文和元数据）")
        else:
            review_strategy = "basic"  # 基础综述
            print("  使用基础综述策略（主要基于元数据）")
        
        # 构建论文内容摘要
        papers_content = ""
        # 记录最终展示顺序中的论文，用于建立编号到论文的映射
        ordered_papers = []
        
        # 先添加有全文的论文（深度分析）
        if papers_with_fulltext:
            papers_content += "\n=== 可获取全文的论文（请进行深度分析） ===\n"
            for i, paper in enumerate(papers_with_fulltext, 1):
                title = paper.get('title') or '未知标题'
                authors = paper.get('authors') or []
                published = paper.get('published') or '未知'
                venue = paper.get('venue') or ''
                source = paper.get('source') or 'unknown'
                citation_count = paper.get('citation_count')
                reference_count = paper.get('reference_count') or 0
                summary = paper.get('summary') or '无摘要'
                entry_id = paper.get('entry_id') or ''
                
                # 获取全文内容
                full_text = paper.get('full_text', {})
                if isinstance(full_text, dict):
                    content_type = full_text.get('content_type', 'metadata_only')
                    content = full_text.get('content')
                    content_source = full_text.get('source', 'unknown')
                else:
                    # 兼容旧格式
                    content_type = 'latex' if paper.get('latex_content') else 'metadata_only'
                    content = paper.get('latex_content')
                    content_source = 'arxiv'
                
                papers_content += f"\n[论文 {i}] {title}\n"
                papers_content += f"作者: {', '.join(authors) if authors else '未知'}\n"
                papers_content += f"发表时间: {published}\n"
                if venue:
                    papers_content += f"来源: {venue} ({source})\n"
                if citation_count is not None:
                    papers_content += f"引用数: {citation_count}, 参考文献数: {reference_count}\n"
                papers_content += f"摘要: {summary}\n"
                
                # 添加全文内容
                if content:
                    if content_type == 'latex':
                        papers_content += f"\n论文LaTeX内容（已解析为可读文本，来源：{content_source}）:\n{content}\n"
                    elif content_type == 'xml':
                        papers_content += f"\n论文XML全文内容（来源：{content_source}）:\n{content[:10000]}...\n"  # 限制长度
                    elif content_type == 'pdf_text':
                        papers_content += f"\n论文PDF文本内容（来源：{content_source}）:\n{content[:10000]}...\n"  # 限制长度
                
                if entry_id:
                    papers_content += f"链接: {entry_id}\n"
                papers_content += "-" * 80 + "\n"
                ordered_papers.append(paper)
        
        # 再添加仅元数据的论文（参考信息）
        if papers_metadata_only:
            papers_content += "\n=== 其他相关论文（基于摘要信息，作为参考） ===\n"
            start_idx = len(papers_with_fulltext) + 1
            for i, paper in enumerate(papers_metadata_only, start=start_idx):
                title = paper.get('title') or '未知标题'
                authors = paper.get('authors') or []
                published = paper.get('published') or '未知'
                venue = paper.get('venue') or ''
                source = paper.get('source') or 'unknown'
                citation_count = paper.get('citation_count')
                reference_count = paper.get('reference_count') or 0
                summary = paper.get('summary') or '无摘要'
                entry_id = paper.get('entry_id') or ''
                
                papers_content += f"\n[论文 {i}] {title}\n"
                papers_content += f"作者: {', '.join(authors) if authors else '未知'}\n"
                papers_content += f"发表时间: {published}\n"
                if venue:
                    papers_content += f"来源: {venue} ({source})\n"
                if citation_count is not None:
                    papers_content += f"引用数: {citation_count}, 参考文献数: {reference_count}\n"
                papers_content += f"摘要: {summary}\n"
                if entry_id:
                    papers_content += f"链接: {entry_id}\n"
                papers_content += "-" * 80 + "\n"
                ordered_papers.append(paper)
        
        # 基于“日期新 / 引用高 / 二者皆是”进行分类，并输出编号列表，供模型在不同段落重点使用
        def _parse_year(p):
            # 优先 published 年份，再次使用 publication_year
            pub = p.get('published') or ''
            year_match = re.search(r'(\d{4})', str(pub))
            if year_match:
                try:
                    return int(year_match.group(1))
                except Exception:
                    pass
            py = p.get('publication_year')
            try:
                return int(py) if py else None
            except Exception:
                return None
        
        current_year = datetime.now().year
        recent_window_years = 2  # 近两年视为“日期新”
        high_citation_threshold = 20  # 引用数阈值（>= 20 视为“引用高”）
        
        recent_indices = []
        high_indices = []
        both_indices = []
        
        for idx, p in enumerate(ordered_papers, 1):
            year = _parse_year(p)
            citations = int(p.get('citation_count', 0) or 0)
            is_recent = (year is not None) and (year >= current_year - recent_window_years)
            is_high = citations >= high_citation_threshold
            if is_recent and is_high:
                both_indices.append(idx)
            elif is_recent:
                recent_indices.append(idx)
            elif is_high:
                high_indices.append(idx)
        
        categorization_instruction = (
            "分类提示（按上面论文编号）：\n"
            f"- 引用高的论文编号：{high_indices if high_indices else '无'}\n"
            f"- 日期新且引用高的论文编号：{both_indices if both_indices else '无'}\n"
            f"- 日期新的论文编号：{recent_indices if recent_indices else '无'}\n\n"
            "写作要求（严格执行）：\n"
            "- 使用“引用高”的论文书写过往被广泛认可的技术发展脉络（回顾部分）。\n"
            "- 使用“日期新且引用高”的论文书写最新先进技术成果（代表性成果与SOTA）。\n"
            "- 使用“日期新”的论文书写最新研究趋势（方向演进、热点与前沿）。\n"
        )
        
        # 构建提示词
        system_prompt = """你是一位专业的学术研究助手，擅长撰写学术综述。请根据提供的论文信息，撰写一篇结构清晰、内容全面的学术综述。

注意：论文的LaTeX内容已经过解析处理，去除了LaTeX命令和结构标记，提取了可读的文本内容、章节结构和主要论述。你可以直接理解这些文本内容，无需担心LaTeX格式问题。

重要要求：
1. 综述必须包含以下四个核心部分（这是综述的主要目的）：
   - 明确背景：通过引言部分描述该研究领域的重要性和背景，说明研究意义和现状
   - 总结技术：对技术进行全面总结，将具有类似性的技术放在一起描述，进行总结性陈述（无需详细刻画技术细节）
   - 提炼趋势：分析当前研究的发展趋势、热点方向和变化特点
   - 展望未来：展望该领域的未来发展方向、潜在挑战和可能的研究机会

2. 综述结构要求：
   - 标题：为综述起一个合适的标题（使用一级标题 # 标题）
   - 摘要：200-300字的摘要，概括综述的主要内容（使用二级标题 ## 摘要）
   - 关键词：3-5个关键词，用分号分隔（使用二级标题 ## 关键词）
   - 引言：明确背景，介绍该研究领域的重要性和背景（使用二级标题 ## 引言）
   - 正文：引言之后必须至少有3-4个二级标题，用于总结技术（将类似技术归类在一起）
   - 趋势分析：分析当前研究的发展趋势、热点方向和变化特点（使用二级标题）
   - 未来展望：展望该领域的未来发展方向、潜在挑战和可能的研究机会（使用二级标题）
   - 参考文献：不需要在综述末尾列出，系统会根据论文语言自动生成相应格式的参考文献列表（中文期刊使用GB/T 7714-2015，英文期刊使用IEEE格式）

2. 标题格式要求（非常重要）：
   - 使用一级标题（# 标题）表示总标题（全文只有一个，居中显示，无数字序号）
   - 使用二级标题（## 标题）表示主要章节：
     * ## 摘要（加粗，无数字序号，标题中不要包含"摘要"以外的任何编号）
     * ## 关键词（加粗，无数字序号，标题中不要包含"关键词"以外的任何编号）
     * ## 引言（加粗，标题中不要包含任何编号，系统会自动添加"0 "编号）
     * ## 其他二级标题（加粗，标题中不要包含任何编号，系统会自动从"1"开始编号：1、2、3...）
   - 使用三级标题（### 标题）表示子章节（仅在引言之后的二级标题下使用，标题中不要包含任何编号，系统会自动添加编号：第一个二级标题下为1.1、1.2...，第二个二级标题下为2.1、2.2...，以此类推）
   - 使用四级标题（#### 标题）表示更细的子章节（仅在三级标题下使用，标题中不要包含任何编号，系统会自动添加编号：如1.1.1、1.1.2...、1.2.1、1.2.2...等）
   - 引言下不需要其他次级标题（不要使用三级或四级标题）
   - 引言之后必须至少有3-4个二级标题，每个二级标题下可以有多个三级标题，每个三级标题下可以有多个四级标题
   - 重要：所有标题中不要包含数字编号（如"1"、"1.1"、"1.1.1"等），系统会自动添加编号

3. 技术总结的写作要求（非常重要）：
   - 将具有类似性的技术放在一起描述，不要分散描述
   - 进行总结性陈述，无需详细刻画技术细节
   - 通常有好几段，每段通过引用多个论文（如[1,2,3]）来描述类似性技术
   - 每段应该概括一类技术方法的特点、优势和应用场景
   - 避免对单个论文进行详细描述，而是对一类技术进行综合总结

4. 段落要求：
   - 段落与段落之间不要有空行（连续段落直接写，不要用空行分隔）
   - 每个段落应该是一个完整的论述单元
   - 使用学术化的语言，确保综述逻辑清晰、条理分明

5. 参考文献引用：
   - 在文中提及论文时，使用[序号]格式标注，如[1]、[2]或[1,2,3]
   - 序号对应论文列表中的顺序（论文列表已标注[论文 1]、[论文 2]等，引用时使用对应序号）
   - 在总结技术时，每段应引用多个相关论文（如[1,2,3]），以体现技术的类似性和共性
   - 不需要在综述末尾列出参考文献，系统会根据论文语言自动生成相应格式的参考文献列表

6. 输出格式示例：
# 综述标题

## 摘要
摘要内容...

## 关键词
关键词1; 关键词2; 关键词3

## 引言
引言段落内容，明确研究背景和重要性[1]。继续论述该领域的研究现状[2,3]...

## 技术方法一：XXX类方法
这类方法的主要特点是...，在...方面有重要应用[4,5,6]。另一类相关研究关注...，提出了...方法[7,8]。这些方法共同的特点是...，适用于...场景[9,10]。

## 技术方法二：YYY类技术
这类技术主要解决...问题，包括...和...等方向[11,12,13]。近年来，研究者提出了...改进方法[14,15]，在...方面取得了显著进展[16,17]。

## 研究趋势
段落内容分析当前发展趋势[18,19]...

## 未来展望
段落内容展望未来发展方向[20,21]...

注意：示例中的标题（如"## 技术方法一"）不包含编号，系统会自动添加编号（引言后的二级标题为1、2、3...，其下的三级标题为1.1、1.2...、2.1、2.2...等）

重要提示：
- 引言之后必须至少有3-4个二级标题用于总结技术（不包括趋势和展望）
- 每个技术总结的二级标题下，应有多段内容，每段引用多个相关论文（如[1,2,3]），描述类似性技术
- 最后必须有"## 研究趋势"和"## 未来展望"两个二级标题
- 技术总结部分无需详细刻画技术细节，而是进行总结性陈述
- 不需要在综述末尾列出参考文献，系统会根据论文语言自动生成相应格式的参考文献列表"""
        
        # 根据综述策略调整提示词
        if review_strategy == "deep":
            strategy_instruction = """
重要提示：你已获得多篇论文的完整全文内容（LaTeX/XML/PDF文本）。请充分利用这些全文内容进行深度分析：
- 深入分析论文的方法、实验设计和结果
- 详细比较不同论文的技术路线和创新点
- 基于全文内容提供更深入的技术洞察和批判性分析
- 重点关注可获取全文的论文，其他论文作为补充参考
"""
        elif review_strategy == "medium":
            strategy_instruction = """
重要提示：你已获得部分论文的完整全文内容。请结合全文和摘要信息：
- 对可获取全文的论文进行深度分析
- 对其他论文基于摘要信息进行概述
- 平衡使用全文内容和摘要信息
"""
        else:
            strategy_instruction = """
重要提示：主要基于论文的标题和摘要信息生成综述。请：
- 充分利用摘要信息进行综合分析
- 基于引用数和发表时间评估论文重要性
- 提供领域概述和趋势分析
"""
        
        user_prompt = f"""请基于以下关于"{keyword}"的论文信息，撰写一篇学术综述。
{strategy_instruction}

{papers_content}

{categorization_instruction}

请用中文撰写综述，严格按照上述格式要求：
1. 包含标题（一级标题）、摘要（二级标题）、关键词（二级标题）、引言（二级标题，明确背景）
2. 引言之后必须至少有3-4个二级标题用于总结技术，将类似技术放在一起描述，每段引用多个论文（如[1,2,3]）
3. 技术总结部分进行总结性陈述，无需详细刻画技术细节
4. 最后必须有"## 研究趋势"和"## 未来展望"两个二级标题
5. 段落之间不要有空行
6. 在文中使用[序号]格式引用论文，技术总结部分每段应引用多个相关论文
7. 不需要在综述末尾列出参考文献，系统会根据论文语言自动生成相应格式的参考文献列表"""
        
        # 调用大模型
        response = client.chat.completions.create(
            model='qwen-max',
            messages=[
                {
                    'role': 'system',
                    'content': system_prompt
                },
                {
                    'role': 'user',
                    'content': user_prompt
                }
            ],
            stream=False,
            temperature=0.7
        )
        
        if response and response.choices and len(response.choices) > 0:
            review = response.choices[0].message.content
            if not review:
                return ("生成综述时出错：返回内容为空。", [])
        else:
            return ("生成综述时出错：API返回格式异常。", [])
        
        # 生成综述后，由大模型选择要引用的文献
        print(f"\n正在由大模型选择要引用的文献（从{len(papers)}篇中选择{final_citation_count}篇）...")
        cited_papers = select_cited_papers_with_llm(papers, review, final_citation_count, keyword)
        
        return (review, cited_papers)
            
    except Exception as e:
        error_msg = f"生成综述时出错: {e}"
        print(error_msg)
        return (error_msg, [])

def review_and_revise_review(review, keyword, papers):
    """
    使用大模型对生成的综述进行审阅和修改
    
    Args:
        review: 生成的综述内容
        keyword: 关键词
        papers: 论文列表
    
    Returns:
        审阅修改后的综述内容
    """
    print(f"\n正在使用大模型审阅和修改综述...")
    
    try:
        # 构建审阅提示词
        system_prompt = """你是一位专业的学术审稿专家，擅长审阅和修改学术综述。请根据以下要求对综述进行审阅和修改：

审阅要求：
1. 标题级别检查：
   - 总标题：一级标题（# 标题），居中，无数字序号
   - 摘要和关键词：二级标题（## 摘要、## 关键词），加粗，无数字序号（标题中不要包含任何编号）
   - 引言：二级标题（## 引言），加粗，标题中不要包含任何编号（系统会自动添加"0 "编号），引言下不需要其他次级标题
   - 引言之后的二级标题：必须至少有3-4个，标题中不要包含任何编号（系统会自动从"1"开始编号：1、2、3...）
   - 三级标题：仅在引言之后的二级标题下使用，标题中不要包含任何编号（系统会自动添加编号：第一个二级标题下为1.1、1.2...，第二个二级标题下为2.1、2.2...，以此类推）
   - 四级标题：仅在三级标题下使用，标题中不要包含任何编号（系统会自动添加编号：如1.1.1、1.1.2...、1.2.1、1.2.2...等）
   - 重要：如果发现标题中包含编号（如"1"、"1.1"、"1.1.1"等），请去除这些编号，只保留标题文本

2. 内容质量检查：
   - 确保综述逻辑清晰、条理分明
   - 确保引用格式正确（使用[序号]格式）
   - 确保段落之间没有空行
   - 确保学术语言规范

3. 结构完整性检查：
   - 必须包含：标题、摘要、关键词、引言（明确背景）
   - 引言之后必须至少有3-4个二级标题用于总结技术（不包括趋势和展望）
   - 技术总结部分应将类似技术放在一起描述，每段引用多个论文
   - 最后必须有"## 研究趋势"和"## 未来展望"两个二级标题
   - 每个二级标题下可以有多个三级标题（视情况而定）

4. 修改原则：
   - 如果发现不符合要求的地方，请直接修改
   - 保持原文的核心内容和观点
   - 优化表达，提高学术规范性
   - 确保所有要求都得到满足

请直接输出修改后的完整综述内容，使用Markdown格式。"""
        
        user_prompt = f"""请审阅以下关于"{keyword}"的学术综述，并根据所有要求进行修改：

{review}

请输出修改后的完整综述内容，确保：
1. 标题级别符合要求（总标题一级居中无序号，摘要关键词二级加粗无序号，引言二级序号0，之后二级从1开始）
2. 引言之后至少有3-4个二级标题用于总结技术（不包括趋势和展望）
3. 技术总结部分将类似技术放在一起描述，每段引用多个论文，进行总结性陈述
4. 最后必须有"## 研究趋势"和"## 未来展望"两个二级标题
5. 段落之间没有空行
6. 引用格式正确（技术总结部分每段应引用多个相关论文）
7. 内容质量高，逻辑清晰"""
        
        # 调用大模型进行审阅
        response = client.chat.completions.create(
            model='qwen-max',
            messages=[
                {
                    'role': 'system',
                    'content': system_prompt
                },
                {
                    'role': 'user',
                    'content': user_prompt
                }
            ],
            stream=False,
            temperature=0.5  # 审阅时使用较低的温度，更保守
        )
        
        if response and response.choices and len(response.choices) > 0:
            revised_review = response.choices[0].message.content
            return revised_review if revised_review else review  # 如果返回为空，返回原综述
        else:
            print("审阅时出错：API返回格式异常，返回原综述")
            return review
            
    except Exception as e:
        error_msg = f"审阅综述时出错: {e}"
        print(error_msg)
        print("返回原综述")
        return review  # 出错时返回原综述

def is_chinese_paper(paper):
    """
    判断论文是否为中文论文
    
    Args:
        paper: 论文信息字典
    
    Returns:
        True if 中文论文, False if 英文论文
    """
    # 检查标题
    title = paper.get('title', '')
    if title:
        # 检查是否包含中文字符
        if re.search(r'[\u4e00-\u9fff]', title):
            return True
    
    # 检查摘要
    summary = paper.get('summary', '')
    if summary:
        if re.search(r'[\u4e00-\u9fff]', summary):
            return True
    
    # 检查venue（期刊/会议名称）
    venue = paper.get('venue', '')
    if venue:
        # 常见中文期刊关键词
        chinese_venue_keywords = ['学报', '期刊', '杂志', '研究', '科学', '工程', '技术', '信息', '计算机', '电子', '通信']
        venue_lower = venue.lower()
        for keyword in chinese_venue_keywords:
            if keyword in venue:
                # 进一步检查是否包含中文字符
                if re.search(r'[\u4e00-\u9fff]', venue):
                    return True
    
    # 默认判断为英文论文
    return False

def format_reference_ieee(paper, index):
    """
    将论文信息格式化为IEEE格式的参考文献
    
    Args:
        paper: 论文信息字典
        index: 参考文献序号
    
    Returns:
        格式化的参考文献字符串
    """
    title = paper.get('title', '未知标题')
    authors = paper.get('authors', [])
    published = paper.get('published', '')
    venue = paper.get('venue', '')
    source = paper.get('source', 'unknown')
    entry_id = paper.get('entry_id', '')
    
    # 格式化作者（IEEE格式：First Last, First Last, and First Last）
    if authors:
        if len(authors) == 1:
            author_str = authors[0]
        elif len(authors) <= 6:
            # 前N-1个用逗号分隔，最后一个用"and"连接
            author_str = ', '.join(authors[:-1]) + ', and ' + authors[-1]
        else:
            # 超过6个作者，只列出前6个，然后加"et al."
            author_str = ', '.join(authors[:6]) + ', et al.'
    else:
        author_str = 'Anonymous'
    
    # 提取年份
    year = ''
    if published:
        year_match = re.search(r'(\d{4})', published)
        if year_match:
            year = year_match.group(1)
    
    # 判断文献类型并格式化（IEEE标准）
    if source.lower() == 'arxiv' or 'arxiv' in venue.lower() or 'arxiv' in title.lower():
        # 预印本格式：[序号] Author(s). "Title," arXiv preprint arXiv:XXXX.XXXX, Year.
        arxiv_id = ''
        if entry_id:
            arxiv_match = re.search(r'arxiv\.org/abs/([\d.]+)', entry_id)
            if arxiv_match:
                arxiv_id = arxiv_match.group(1)
        if arxiv_id:
            ref_str = f"[{index}] {author_str}, \"{title},\" arXiv preprint arXiv:{arxiv_id}"
        else:
            ref_str = f"[{index}] {author_str}, \"{title},\" arXiv preprint"
        if year:
            ref_str += f", {year}"
        ref_str += "."
    elif venue:
        # 判断是期刊还是会议
        venue_lower = venue.lower()
        if any(keyword in venue_lower for keyword in ['conference', 'proceedings', 'workshop', 'symposium', 'meeting']):
            # 会议论文格式：[序号] Author(s), "Title," in Proc. Conference Name, Location, Year, pp. pages.
            ref_str = f"[{index}] {author_str}, \"{title},\" in Proc. {venue}"
            if year:
                ref_str += f", {year}"
            ref_str += "."
        else:
            # 期刊论文格式：[序号] Author(s), "Title," Journal Name, vol. vol, no. no, pp. pages, Month Year.
            ref_str = f"[{index}] {author_str}, \"{title},\" {venue}"
            if year:
                ref_str += f", {year}"
            ref_str += "."
    else:
        # 其他类型
        ref_str = f"[{index}] {author_str}, \"{title}\""
        if year:
            ref_str += f", {year}"
        if venue:
            ref_str += f", {venue}"
        ref_str += "."
    
    return ref_str

def format_reference_gb7714(paper, index):
    """
    将论文信息格式化为GB/T 7714-2015格式的参考文献
    
    Args:
        paper: 论文信息字典
        index: 参考文献序号
    
    Returns:
        格式化的参考文献字符串
    """
    title = paper.get('title', '未知标题')
    authors = paper.get('authors', [])
    published = paper.get('published', '')
    venue = paper.get('venue', '')
    source = paper.get('source', 'unknown')
    entry_id = paper.get('entry_id', '')
    
    # 格式化作者（最多3个，超过用"等"）
    if authors:
        if len(authors) <= 3:
            author_str = ', '.join(authors)
        else:
            author_str = ', '.join(authors[:3]) + ', 等'
    else:
        author_str = '佚名'
    
    # 提取年份
    year = ''
    if published:
        # 尝试从日期中提取年份
        year_match = re.search(r'(\d{4})', published)
        if year_match:
            year = year_match.group(1)
    
    # 判断文献类型并格式化（GB/T 7714-2015标准）
    if source.lower() == 'arxiv' or 'arxiv' in venue.lower() or 'arxiv' in title.lower():
        # 预印本格式：[序号] 作者. 题名[EB/OL]. (发表日期)[引用日期]. 获取和访问路径.
        date_str = published if published else '未知日期'
        ref_str = f"[{index}] {author_str}. {title}[EB/OL]. ({date_str})"
        if entry_id:
            ref_str += f"[{datetime.now().strftime('%Y-%m-%d')}]. {entry_id}."
        else:
            ref_str += "."
    elif venue:
        # 判断是期刊还是会议
        venue_lower = venue.lower()
        if any(keyword in venue_lower for keyword in ['conference', 'proceedings', 'workshop', 'symposium', 'meeting']):
            # 会议论文格式：[序号] 作者. 题名[C]//会议名称. 会议地点: 出版者, 出版年: 起止页码.
            ref_str = f"[{index}] {author_str}. {title}[C]//{venue}"
            if year:
                ref_str += f". {year}"
            ref_str += "."
        else:
            # 期刊论文格式：[序号] 作者. 题名[J]. 刊名, 出版年, 卷(期): 起止页码.
            ref_str = f"[{index}] {author_str}. {title}[J]. {venue}"
            if year:
                ref_str += f", {year}"
            ref_str += "."
    else:
        # 其他类型，使用通用格式：[序号] 作者. 题名[M/D/R等]. 出版地: 出版者, 出版年.
        ref_str = f"[{index}] {author_str}. {title}"
        if year:
            ref_str += f". {year}"
        if venue:
            ref_str += f". {venue}"
        ref_str += "."
    
    return ref_str

def save_review_to_docx(keyword, papers, unique_final_papers, review, filename):
    """
    将综述保存为格式化的docx文件
    
    Args:
        keyword: 关键词
        papers: 所有搜索到的论文
        unique_final_papers: 最终选择的论文
        review: 生成的综述内容
        filename: 保存的文件名
    """
    doc = Document()
    
    # 设置字体函数：中文使用仿宋，英文、数字使用Times New Roman
    def set_font(run, chinese_font='仿宋', english_font='Times New Roman', size=12):
        """
        设置字体：中文使用仿宋，英文、数字使用Times New Roman
        """
        # 设置英文字体
        run.font.name = english_font
        # 设置中文字体（东亚字体）
        run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
        run.font.size = Pt(size)
    
    # 设置文本字体：中文用仿宋，数字和字母用Times New Roman
    def set_text_font(para, text, chinese_font='仿宋', english_font='Times New Roman', size=12):
        """
        将文本按字符分割，中文用仿宋，数字和字母用Times New Roman
        """
        # 匹配中文字符、数字字母、标点符号
        pattern = r'([\u4e00-\u9fff]+|[a-zA-Z0-9]+|[^\u4e00-\u9fff\w\s]+|\s+)'
        parts = re.findall(pattern, text)
        
        for part in parts:
            if not part.strip():
                # 空白字符
                run = para.add_run(part)
                set_font(run, chinese_font=chinese_font, english_font=english_font, size=size)
            elif re.match(r'^[a-zA-Z0-9]+$', part):
                # 纯数字和字母，使用Times New Roman
                run = para.add_run(part)
                run.font.name = english_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)  # 保持中文字体设置
                run.font.size = Pt(size)
            elif re.match(r'^[\u4e00-\u9fff]+$', part):
                # 纯中文，使用仿宋
                run = para.add_run(part)
                set_font(run, chinese_font=chinese_font, english_font=english_font, size=size)
            else:
                # 混合内容，需要逐字符处理
                for char in part:
                    if re.match(r'^[a-zA-Z0-9]$', char):
                        # 数字和字母
                        run = para.add_run(char)
                        run.font.name = english_font
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                        run.font.size = Pt(size)
                    else:
                        # 中文或其他字符
                        run = para.add_run(char)
                        set_font(run, chinese_font=chinese_font, english_font=english_font, size=size)
    
    # 兼容旧函数名
    def set_chinese_font(run, font_name='仿宋', size=12):
        set_font(run, chinese_font=font_name, size=size)
    
    # 标题（仿宋小三，15磅）
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_text_font(title_para, f'学术综述：{keyword}', chinese_font='仿宋', english_font='Times New Roman', size=15)
    for run in title_para.runs:
        run.font.bold = True  # 加粗
    
    # 添加空行
    doc.add_paragraph()
    
    # 基本信息（一级标题格式：1.）
    info_heading_para = doc.add_paragraph()
    set_text_font(info_heading_para, '1. 基本信息', chinese_font='仿宋', english_font='Times New Roman', size=14)
    for run in info_heading_para.runs:
        run.font.bold = True
    info_heading_para.paragraph_format.space_after = Pt(12)
    info_heading_para.paragraph_format.space_before = Pt(12)
    
    info_para = doc.add_paragraph()
    # 添加"关键词："部分
    run1 = info_para.add_run('关键词：')
    run1.bold = True
    set_font(run1, chinese_font='仿宋', english_font='Times New Roman', size=12)
    # 添加关键词内容（使用set_text_font处理数字和字母）
    for char in keyword:
        if re.match(r'^[a-zA-Z0-9]$', char):
            # 数字和字母用Times New Roman
            run = info_para.add_run(char)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            run.font.size = Pt(12)
        else:
            # 中文或其他字符用仿宋
            run = info_para.add_run(char)
            set_font(run, chinese_font='仿宋', english_font='Times New Roman', size=12)
    
    info_para = doc.add_paragraph()
    run1 = info_para.add_run('搜索到的论文数量：')
    run1.bold = True
    set_font(run1, chinese_font='仿宋', english_font='Times New Roman', size=12)
    set_text_font(info_para, f'{len(papers)} 篇（去重后）', chinese_font='仿宋', english_font='Times New Roman', size=12)
    
    info_para = doc.add_paragraph()
    run1 = info_para.add_run('最终选择的论文数量：')
    run1.bold = True
    set_font(run1, chinese_font='仿宋', english_font='Times New Roman', size=12)
    set_text_font(info_para, f'{len(unique_final_papers)} 篇', chinese_font='仿宋', english_font='Times New Roman', size=12)
    
    # 统计来源
    sources = {}
    for paper in unique_final_papers:
        source = paper.get('source', 'unknown')
        sources[source] = sources.get(source, 0) + 1
    if sources:
        info_para = doc.add_paragraph()
        run1 = info_para.add_run('数据来源：')
        run1.bold = True
        set_font(run1, chinese_font='仿宋', english_font='Times New Roman', size=12)
        set_text_font(info_para, ', '.join([f'{k}({v}篇)' for k, v in sources.items()]), chinese_font='仿宋', english_font='Times New Roman', size=12)
    
    # 添加生成时间
    info_para = doc.add_paragraph()
    run1 = info_para.add_run('生成时间：')
    run1.bold = True
    set_font(run1, chinese_font='仿宋', english_font='Times New Roman', size=12)
    set_text_font(info_para, datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'), chinese_font='仿宋', english_font='Times New Roman', size=12)
    
    # 添加分页符
    doc.add_page_break()
    
    # 处理综述内容，解析Markdown格式的标题和段落
    review_lines = review.split('\n')
    
    # 查找"参考文献"部分的位置，如果存在则跳过（我们会在文档末尾单独生成）
    ref_section_start = -1
    for idx, line in enumerate(review_lines):
        stripped = line.strip()
        if stripped.startswith('##') and ('参考文献' in stripped or '参考' in stripped and '文献' in stripped):
            ref_section_start = idx
            break
    
    # 如果找到参考文献部分，只处理之前的内容
    if ref_section_start >= 0:
        review_lines = review_lines[:ref_section_start]
    
    # 标题编号计数器
    level1_counter = 0  # 引言后的二级标题编号（1, 2, 3...）
    level2_counter = 0  # 三级标题编号（在每个二级标题下从1开始：1.1, 1.2... 2.1, 2.2...）
    level3_counter = 0  # 四级标题编号（如果需要：1.1.1, 1.1.2...）
    last_level = 0  # 上一个标题的级别
    numbering_started = False  # 标记是否开始编号（从引言开始）
    after_intro = False  # 标记是否在引言之后（引言下不需要其他次级标题）
    first_level1_title = True  # 标记是否是第一个一级标题（总标题，居中无序号）
    
    # 定义不需要编号的标题关键词
    no_numbering_keywords = ['摘要', '关键词']
    
    # 辅助函数：去除标题中已有的编号
    def remove_existing_numbering(text):
        """去除标题开头的编号（如"1 "、"2 "、"1.1 "、"1.1.1 "、"0 "等）"""
        import re
        # 匹配开头的编号模式：
        # - "0 "（引言）
        # - "1 "、"2 "、"3 "等（引言后的二级标题）
        # - "1.1 "、"1.2 "、"2.1 "等（三级标题）
        # - "1.1.1 "、"1.1.2 "等（四级标题）
        # 确保编号后必须有空格，避免误匹配标题中的数字
        pattern = r'^(0\s+|\d+(\.\d+)*(\.\d+)?\s+)'
        return re.sub(pattern, '', text).strip()
    
    i = 0
    while i < len(review_lines):
        line = review_lines[i].strip()
        
        # 跳过空行（段落之间不要空行）
        if not line:
            i += 1
            continue
        
        # 检查是否是标题
        if line.startswith('#### '):
            # 四级标题（仿宋小四，12磅，格式：1.1.1, 1.1.2... 1.2.1, 1.2.2...）
            title_text = line[5:].strip()
            # 去除标题中可能已有的编号
            title_text = remove_existing_numbering(title_text)
            
            # 如果还没有开始编号（引言之前），则不编号
            if not numbering_started:
                numbered_title = title_text
            # 如果是在引言之后（引言下不需要其他次级标题），跳过
            elif after_intro:
                i += 1
                continue
            else:
                # 引言之后的二级标题下的三级标题下的四级标题，格式：1.1.1, 1.1.2... 1.2.1, 1.2.2...
                # 更新计数器：如果上一个标题是四级，则递增；否则重置为1
                if last_level == 4:
                    level3_counter += 1
                else:
                    level3_counter = 1
                
                # 生成编号：{level1_counter}.{level2_counter}.{level3_counter}
                number = f"{level1_counter}.{level2_counter}.{level3_counter} "
                numbered_title = number + title_text
            
            heading_para = doc.add_paragraph()
            set_text_font(heading_para, numbered_title, chinese_font='仿宋', english_font='Times New Roman', size=12)
            for run in heading_para.runs:
                run.font.bold = True
            heading_para.paragraph_format.space_after = Pt(12)
            heading_para.paragraph_format.space_before = Pt(12)
            
            last_level = 4
        elif line.startswith('### '):
            # 三级标题（仿宋小四，12磅，格式：1.1, 1.2... 2.1, 2.2...）
            title_text = line[4:].strip()
            # 去除标题中可能已有的编号
            title_text = remove_existing_numbering(title_text)
            
            # 如果还没有开始编号（引言之前），则不编号
            if not numbering_started:
                numbered_title = title_text
            # 如果是在引言之后（引言下不需要其他次级标题），跳过
            elif after_intro:
                i += 1
                continue
            else:
                # 引言之后的二级标题下的三级标题，格式：1.1, 1.2... 2.1, 2.2...
                # 更新计数器：如果上一个标题是三级，则递增；否则重置为1
                if last_level == 3:
                    level2_counter += 1
                else:
                    level2_counter = 1
                level3_counter = 0  # 重置四级标题计数器
                
                # 生成编号：{level1_counter}.{level2_counter}
                number = f"{level1_counter}.{level2_counter} "
                numbered_title = number + title_text
            
            heading_para = doc.add_paragraph()
            set_text_font(heading_para, numbered_title, chinese_font='仿宋', english_font='Times New Roman', size=12)
            for run in heading_para.runs:
                run.font.bold = True
            heading_para.paragraph_format.space_after = Pt(12)
            heading_para.paragraph_format.space_before = Pt(12)
            
            last_level = 3
        elif line.startswith('## '):
            # 二级标题（仿宋小四，12磅）
            title_text = line[3:].strip()
            # 去除标题中可能已有的编号（但保留"0 "用于引言的特殊处理）
            # 先检查是否是引言，如果是引言且已有"0 "，则保留；否则去除编号
            if '引言' not in title_text:
                title_text = remove_existing_numbering(title_text)
            elif title_text.startswith('0 '):
                # 引言已有"0 "编号，去除它（后面会重新添加）
                title_text = title_text[2:].strip()
            
            # 检查是否是引言（开始编号）
            if '引言' in title_text and not numbering_started:
                numbering_started = True
                after_intro = True  # 标记在引言之后（引言下不需要其他次级标题）
                level1_counter = 0  # 引言后的二级标题从1开始（会在遇到第一个二级标题时设置为1）
                level2_counter = 0  # 三级标题计数器
                level3_counter = 0
                number = "0 "  # 引言编号为"0 "（不带点）
                numbered_title = number + title_text
                last_level = 2  # 引言是二级标题
            # 检查是否是不需要编号的标题（摘要、关键词）
            elif any(keyword in title_text for keyword in no_numbering_keywords):
                numbered_title = title_text
                last_level = 2
            # 如果还没有开始编号（引言之前），则不编号
            elif not numbering_started:
                numbered_title = title_text
                last_level = 2
            # 如果是在引言之后，遇到第一个二级标题，清除after_intro标志，从1开始编号
            elif after_intro:
                after_intro = False  # 清除引言后标志
                level1_counter = 1  # 第一个二级标题从1开始
                level2_counter = 0  # 重置三级标题计数器
                level3_counter = 0  # 重置四级标题计数器
                number = f"{level1_counter} "
                numbered_title = number + title_text
                last_level = 2
            else:
                # 引言之后的后续二级标题，递增编号（2, 3, 4...）
                level1_counter += 1
                level2_counter = 0  # 重置三级标题计数器
                level3_counter = 0  # 重置四级标题计数器
                
                # 生成编号（格式：2, 3, 4等）
                number = f"{level1_counter} "
                numbered_title = number + title_text
                last_level = 2
            
            heading_para = doc.add_paragraph()
            set_text_font(heading_para, numbered_title, chinese_font='仿宋', english_font='Times New Roman', size=12)
            for run in heading_para.runs:
                run.font.bold = True
            heading_para.paragraph_format.space_after = Pt(12)
            heading_para.paragraph_format.space_before = Pt(12)
        elif line.startswith('# '):
            # 一级标题（仿宋四号，14磅）
            title_text = line[2:].strip()
            
            # 第一个一级标题是总标题，居中，不需要数字序号
            if first_level1_title:
                numbered_title = title_text
                heading_para = doc.add_paragraph()
                heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
                set_text_font(heading_para, numbered_title, chinese_font='仿宋', english_font='Times New Roman', size=14)
                for run in heading_para.runs:
                    run.font.bold = True
                heading_para.paragraph_format.space_after = Pt(12)
                heading_para.paragraph_format.space_before = Pt(12)
                first_level1_title = False
                last_level = 1
            # 其他一级标题按1. 2. 等编号（但根据新要求，一级标题应该很少出现）
            else:
                # 检查是否是不需要编号的标题
                if any(keyword in title_text for keyword in no_numbering_keywords):
                    numbered_title = title_text
                else:
                    # 更新计数器
                    level1_counter += 1
                    level2_counter = 0  # 重置二级标题计数器
                    level3_counter = 0  # 重置三级标题计数器
                    
                    # 生成编号
                    number = f"{level1_counter}. "
                    numbered_title = number + title_text
                
                heading_para = doc.add_paragraph()
                set_text_font(heading_para, numbered_title, chinese_font='仿宋', english_font='Times New Roman', size=14)
                for run in heading_para.runs:
                    run.font.bold = True
                heading_para.paragraph_format.space_after = Pt(12)
                heading_para.paragraph_format.space_before = Pt(12)
                last_level = 1
        else:
            # 普通段落（中文用仿宋，数字和字母用Times New Roman）
            current_para = doc.add_paragraph()
            
            # 处理文本中的参考文献引用 [1] 或 [1,2,3]，转换为上标
            text = line
            
            # 去掉类似*的不必要标点符号（保留必要的标点符号）
            # 去掉单独的*号（前后有空格或开头结尾的*）
            text = re.sub(r'\s*\*\s*', ' ', text)  # 去掉单独的*号及其周围空格
            text = re.sub(r'^\*+', '', text)  # 去掉开头的*号
            text = re.sub(r'\*+$', '', text)  # 去掉结尾的*号
            text = re.sub(r'\*{2,}', '', text)  # 去掉连续的多个*号
            # 清理多余的空格
            text = re.sub(r'\s+', ' ', text).strip()
            
            # 使用正则表达式匹配 [数字] 或 [数字,数字,...]
            ref_pattern = r'\[(\d+(?:,\d+)*)\]'
            
            last_end = 0
            for match in re.finditer(ref_pattern, text):
                # 添加引用前的文本（使用set_text_font处理数字和字母）
                if match.start() > last_end:
                    prefix_text = text[last_end:match.start()]
                    set_text_font(current_para, prefix_text, chinese_font='仿宋', english_font='Times New Roman', size=12)
                
                # 添加引用（上标格式，数字用Times New Roman）
                ref_text = '[' + match.group(1) + ']'
                # 引用中的数字和括号需要分别处理
                for char in ref_text:
                    run = current_para.add_run(char)
                    if char.isdigit():
                        # 数字用Times New Roman
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                    else:
                        # 括号用仿宋
                        set_font(run, chinese_font='仿宋', english_font='Times New Roman', size=12)
                    run.font.size = Pt(12)
                    run.font.superscript = True  # 设置为上标
                
                last_end = match.end()
            
            # 添加剩余文本（使用set_text_font处理数字和字母）
            if last_end < len(text):
                suffix_text = text[last_end:]
                set_text_font(current_para, suffix_text, chinese_font='仿宋', english_font='Times New Roman', size=12)
            
            # 如果段落为空（只有引用），至少添加一个run
            if not current_para.runs:
                set_text_font(current_para, text, chinese_font='仿宋', english_font='Times New Roman', size=12)
            
            # 设置段落格式
            current_para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2个字符（约24磅）
            current_para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            current_para.paragraph_format.space_after = Pt(0)  # 段间距0磅（段落之间不要空行）
        
        i += 1
    
    # 添加分页符
    doc.add_page_break()
    
    # 参考文献列表（GB/T 7714-2015格式）
    ref_heading_para = doc.add_paragraph()
    ref_heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_text_font(ref_heading_para, '参考文献', chinese_font='仿宋', english_font='Times New Roman', size=14)
    for run in ref_heading_para.runs:
        run.font.bold = True  # 加粗
    ref_heading_para.paragraph_format.space_after = Pt(12)
    
    # 根据论文语言选择引用格式：中文期刊使用GB/T 7714-2015，英文期刊使用IEEE格式
    for i, paper in enumerate(unique_final_papers, 1):
        # 判断论文语言
        if is_chinese_paper(paper):
            ref_text = format_reference_gb7714(paper, i)
            # GB/T 7714-2015格式：悬挂缩进
            ref_para = doc.add_paragraph()
            set_text_font(ref_para, ref_text, chinese_font='仿宋', english_font='Times New Roman', size=12)
            ref_para.paragraph_format.left_indent = Pt(24)  # 整个段落左缩进
            ref_para.paragraph_format.first_line_indent = Pt(-24)  # 首行悬挂缩进（负值表示向左）
            ref_para.paragraph_format.space_after = Pt(6)  # 段后间距
            ref_para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
        else:
            ref_text = format_reference_ieee(paper, i)
            # IEEE格式：左对齐，首行缩进
            ref_para = doc.add_paragraph()
            set_text_font(ref_para, ref_text, chinese_font='仿宋', english_font='Times New Roman', size=12)
            ref_para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
            ref_para.paragraph_format.space_after = Pt(6)  # 段后间距
            ref_para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    
    # 保存文档
    import os
    import time
    
    # 检查文件是否已存在
    if os.path.exists(filename):
        print(f"\n警告：文件 '{filename}' 已存在")
        print("可能的原因：")
        print("  1. 文件正在被其他程序（如 Word）打开")
        print("  2. 之前已生成过同名文件")
        print("\n请关闭可能打开该文件的程序，然后重试")
        
        # 尝试生成新文件名
        base_name = os.path.splitext(filename)[0]
        extension = os.path.splitext(filename)[1]
        counter = 1
        new_filename = f"{base_name}_{counter}{extension}"
        while os.path.exists(new_filename):
            counter += 1
            new_filename = f"{base_name}_{counter}{extension}"
        
        print(f"\n或者使用新文件名保存: {new_filename}")
        use_new = input("是否使用新文件名保存？(y/n，默认y): ").strip().lower()
        if use_new in ['', 'y', 'yes', '是']:
            filename = new_filename
        else:
            print("取消保存操作")
            return
    
    try:
        # 尝试保存文件
        doc.save(filename)
        print(f"\n✓ 综述已成功保存到文件: {filename}")
        print(f"完整路径: {os.path.abspath(filename)}")
    except PermissionError as e:
        print(f"\n✗ 保存失败：权限错误")
        print(f"错误信息: {e}")
        print("\n可能的原因：")
        print("  1. 文件正在被其他程序（如 Microsoft Word）打开")
        print("  2. 文件被设置为只读")
        print("  3. 没有写入权限")
        print("\n解决方案：")
        print("  1. 关闭可能打开该文件的所有程序")
        print("  2. 检查文件属性，确保不是只读")
        print("  3. 检查文件夹的写入权限")
        print("  4. 尝试使用不同的文件名")
        
        # 提供重试选项
        retry = input("\n是否尝试使用新文件名保存？(y/n): ").strip().lower()
        if retry in ['y', 'yes', '是']:
            base_name = os.path.splitext(filename)[0]
            extension = os.path.splitext(filename)[1]
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            new_filename = f"{base_name}_{timestamp}{extension}"
            try:
                doc.save(new_filename)
                print(f"\n✓ 综述已成功保存到新文件: {new_filename}")
                print(f"完整路径: {os.path.abspath(new_filename)}")
            except Exception as e2:
                print(f"\n✗ 使用新文件名保存也失败: {e2}")
    except Exception as e:
        print(f"\n✗ 保存文件时发生错误: {e}")
        print(f"错误类型: {type(e).__name__}")
        print("\n请检查：")
        print("  1. 文件路径是否有效")
        print("  2. 是否有足够的磁盘空间")
        print("  3. 文件夹是否存在")

